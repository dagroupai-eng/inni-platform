import streamlit as st
import os
import re
from dotenv import load_dotenv
from file_analyzer import UniversalFileAnalyzer
from dspy_analyzer import EnhancedArchAnalyzer
from prompt_processor import load_blocks, load_custom_blocks
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# 환경변수 로드 (안전하게 처리)
try:
    load_dotenv()
except UnicodeDecodeError:
    # .env 파일에 인코딩 문제가 있는 경우 무시
    pass

# 페이지 설정
st.set_page_config(
    page_title="도시 프로젝트 분석",
    page_icon=None,
    layout="wide"
)

# 제목
st.title("도시 프로젝트 분석")
st.markdown("**도시 프로젝트 문서 분석 (PDF, Excel, CSV, 텍스트, JSON 지원)**")

# Session state 초기화
if 'project_name' not in st.session_state:
    st.session_state.project_name = ""
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'selected_blocks' not in st.session_state:
    st.session_state.selected_blocks = []
if 'pdf_text' not in st.session_state:
    st.session_state.pdf_text = ""
if 'pdf_uploaded' not in st.session_state:
    st.session_state.pdf_uploaded = False

# 블록들을 JSON 파일에서 로드
def get_example_blocks():
    """blocks.json에서 예시 블록들을 로드합니다."""
    return load_blocks()

def create_word_document(project_name, analysis_results):
    """분석 결과를 Word 문서로 생성합니다."""
    doc = Document()
    
    # 제목
    doc.add_heading(f'건축 프로젝트 분석 보고서: {project_name}', 0)
    
    # 각 분석 결과 추가
    for block_id, result in analysis_results.items():
        # 블록 이름 찾기
        block_name = "사용자 정의 블록"
        if block_id.startswith('custom_'):
            custom_blocks = load_custom_blocks()
            for block in custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        else:
            example_blocks = get_example_blocks()
            for block in example_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        
        # 섹션 제목
        doc.add_heading(block_name, level=1)
        
        # Word 표 형식으로 처리
        add_content_with_tables(doc, result)
        doc.add_paragraph()  # 빈 줄
    
    return doc

def add_content_with_tables(doc, text):
    """텍스트를 분석하여 표는 Word 표로, 일반 텍스트는 문단으로 추가합니다."""
    import re
    
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 표 시작 패턴 확인 (개선된 방식)
        if is_table_line(line):
            # 표 데이터 수집
            table_lines = [line]
            i += 1
            
            # 연속된 표 줄들 수집 (개선된 방식)
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Word 표 생성
            create_word_table(doc, table_lines)
            continue
        
        # 일반 텍스트 처리
        if line:
            # Markdown 헤더 처리
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                doc.add_heading(header_text, level=min(level, 6))
            else:
                # 리스트 처리
                if line.startswith('- '):
                    line = '• ' + line[2:]
                elif line.startswith('* '):
                    line = '• ' + line[2:]
                
                # 볼드 텍스트 처리 (**text**)
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                
                doc.add_paragraph(line)
        
        i += 1

def create_word_table(doc, table_lines):
    """Markdown 표 줄들을 Word 표로 변환합니다."""
    if not table_lines:
        return
    
    # 표 데이터 파싱
    table_data = []
    for line in table_lines:
        # |로 구분된 셀들 추출
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # 첫 번째와 마지막 빈 요소 제거
        if cells:
            table_data.append(cells)
    
    if not table_data:
        return
    
    # 첫 번째 행이 헤더 구분선인지 확인 (--- 형태)
    if len(table_data) > 1 and all(cell == '---' or cell == '------' or cell == '' for cell in table_data[1]):
        headers = table_data[0]
        data_rows = table_data[2:]
    else:
        headers = None
        data_rows = table_data
    
    # 열 수 결정
    max_cols = max(len(row) for row in table_data) if table_data else 2
    
    # Word 표 생성 - 개선된 방식
    try:
        table = doc.add_table(rows=len(data_rows) + (1 if headers else 0), cols=max_cols)
        table.style = 'Table Grid'
        
        # 표 자동 크기 조절 활성화
        table.allow_autofit = True
        table.autofit = True
        
        # 헤더 추가
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                if i < len(header_row.cells):
                    cell = header_row.cells[i]
                    cell.text = clean_text_for_pdf(header)
                    
                    # 헤더 스타일링 강화
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                        # 셀 패딩 조정
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(2)
        
        # 데이터 행 추가
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data_rows):
            if start_row + i < len(table.rows):
                table_row = table.rows[start_row + i]
                for j, cell_data in enumerate(row_data):
                    if j < len(table_row.cells):
                        cell = table_row.cells[j]
                        cell.text = clean_text_for_pdf(cell_data)
                        
                        # 셀 스타일링
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                            # 셀 패딩 조정
                            paragraph.paragraph_format.space_before = Pt(1)
                            paragraph.paragraph_format.space_after = Pt(1)
        
        # 표 후 빈 줄 추가
        doc.add_paragraph()
        
    except Exception as e:
        print(f"Word 표 생성 오류: {e}")
        # 오류 발생 시 텍스트로 대체
        doc.add_paragraph("[표 생성 실패 - 원본 데이터]")
        for row in table_data:
            doc.add_paragraph(" | ".join(row))
        doc.add_paragraph()

def is_table_line(line):
    """한 줄이 표 행인지 확인"""
    if not line:
        return False
    
    # | 구분자가 있는 경우 (마크다운 표 형식)
    if '|' in line and line.count('|') >= 2:
        return True
    
    # 탭으로 구분된 경우
    if '\t' in line:
        return True
    
    # 2개 이상의 공백으로 구분된 경우 (정렬된 텍스트)
    if re.search(r'\s{2,}', line):
        return True
    
    return False

def is_table_format(text):
    """텍스트가 표 형식인지 확인"""
    try:
        if not text or not isinstance(text, str):
            return False
            
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False
        
        # 1. 마크다운 표 형식 확인 (| 구분자)
        pipe_count = text.count('|')
        if pipe_count >= 3:  # 최소 1x2 표를 위해서는 3개의 | 필요
            # 구분선이 있는지 확인 (표의 특징)
            for line in lines:
                line = line.strip()
                if re.match(r'^[\s\-=_:|]+\s*$', line):
                    return True
            # 구분선이 없어도 |가 많이 있으면 표로 간주
            if pipe_count >= 6:
                return True
        
        # 2. 구분선 확인 (마크다운 표 구분선)
        for line in lines:
            line = line.strip()
            if re.match(r'^[\s\-=_:|]+\s*$', line):
                return True
        
        # 3. 탭 구분자 확인
        tab_count = sum(1 for line in lines if '\t' in line)
        if tab_count >= 2:
            return True
        
        return False
        
    except Exception as e:
        print(f"표 형식 확인 오류: {e}")
        return False

def clean_text_for_pdf(text):
    """PDF/Word용 텍스트 정리"""
    if not text:
        return ""
    
    import re
    
    # HTML 태그 제거
    text = re.sub(r'<[^>]+>', '', text)
    
    # Markdown 볼드 제거 (**text** -> text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Markdown 이탤릭 제거 (*text* -> text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # 특수 문자 정리
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    
    # 연속된 공백 정리
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


# 사이드바 - 설정
with st.sidebar:
    
    st.header("설정")
    
    # Streamlit secrets와 환경변수 모두 확인
    from dotenv import load_dotenv
    load_dotenv()
    
    # Streamlit secrets에서 먼저 확인
    api_key = st.secrets.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY")
    
    if not api_key:
        st.error("ANTHROPIC_API_KEY가 설정되지 않았습니다!")
        st.info("다음 중 하나의 방법으로 API 키를 설정해주세요:")
        st.code("""
# 방법 1: .streamlit/secrets.toml 파일에 추가
[secrets]
ANTHROPIC_API_KEY = "your_api_key_here"

# 방법 2: .env 파일에 추가
ANTHROPIC_API_KEY=your_api_key_here
        """, language="toml")
        st.stop()
    else:
        st.success("API 키가 설정되었습니다!")
        st.info(f"API 키 길이: {len(api_key)}자")
        st.info(f"키 소스: {'Streamlit Secrets' if st.secrets.get('ANTHROPIC_API_KEY') else '환경변수'}")

# 메인 컨텐츠
tab1, tab2, tab3, tab4 = st.tabs(["기본 정보 & 파일 업로드", "분석 블록 선택", "분석 실행", "결과 다운로드"])

with tab1:
    st.header("프로젝트 기본 정보")
    
    # 프로젝트명 입력
    project_name = st.text_input(
        "프로젝트명", 
        placeholder="예: 서울 도심 재생 프로젝트", 
        help="도시 프로젝트의 이름을 입력하세요"
    )
    
    # 기본 정보 입력 섹션
    st.subheader("프로젝트 개요")
    
    location = st.text_input(
        "위치/지역",
        placeholder="예: 서울시 중구 명동 일대",
        help="프로젝트가 진행될 도시 지역을 입력하세요"
    )
    
    # 프로젝트 목표
    project_goals = st.text_area(
        "프로젝트 목표",
        placeholder="프로젝트의 목표, 비전, 기대효과를 입력하세요...",
        height=80,
        help="도시 프로젝트의 목표, 비전, 기대효과를 입력하세요"
    )
    
    # 추가 정보
    st.subheader("추가 정보")
    additional_info = st.text_area(
        "추가 정보",
        placeholder="프로젝트와 관련된 특별한 요구사항, 제약조건, 참고사항 등을 입력하세요...",
        height=80,
        help="프로젝트와 관련된 특별한 요구사항, 제약조건, 참고사항 등을 자유롭게 입력하세요"
    )
    
    st.markdown("---")
    st.header("파일 업로드")
    
    uploaded_file = st.file_uploader(
        "파일을 업로드하세요",
        type=['pdf', 'xlsx', 'xls', 'csv', 'txt', 'json'],
        help="도시 프로젝트 관련 문서를 업로드하세요 (PDF, Excel, CSV, 텍스트, JSON 지원)"
    )
    
    if uploaded_file is not None:
        st.success(f"파일 업로드 완료: {uploaded_file.name}")
        
        # 파일 확장자 확인
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # 메모리에서 직접 파일 분석 (임시 파일 생성 없음)
        file_analyzer = UniversalFileAnalyzer()
        
        # 파일 분석 (메모리 기반)
        with st.spinner(f"{file_extension.upper()} 파일 분석 중..."):
            analysis_result = file_analyzer.analyze_file_from_bytes(
                uploaded_file.getvalue(), 
                file_extension, 
                uploaded_file.name
            )
            
        if analysis_result['success']:
            st.success(f"{file_extension.upper()} 파일 분석 완료!")
            
            # 파일 정보 표시 (파일 크기는 업로드된 파일에서 직접 계산)
            file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
            st.info(f"파일 정보: {file_size_mb:.2f}MB, {analysis_result['word_count']}단어, {analysis_result['char_count']}문자")
            
            # 파일 형식별 특별 정보 표시
            if analysis_result['file_type'] == 'excel':
                st.info(f"Excel 시트: {', '.join(analysis_result['sheet_names'])} ({analysis_result['sheet_count']}개 시트)")
            elif analysis_result['file_type'] == 'csv':
                st.info(f"CSV 데이터: {analysis_result['shape'][0]}행 × {analysis_result['shape'][1]}열")
            
            # 세션에 저장
            st.session_state['pdf_text'] = analysis_result['text']  # 기존 변수명 유지
            st.session_state['pdf_uploaded'] = True
            st.session_state['file_type'] = analysis_result['file_type']
            st.session_state['file_analysis'] = analysis_result
            st.session_state['uploaded_file'] = uploaded_file  # 파일 객체 저장
            
            # 텍스트 미리보기
            with st.expander(f"{file_extension.upper()} 내용 미리보기"):
                st.text(analysis_result['preview'])
        else:
            st.error(f"{file_extension.upper()} 파일 분석에 실패했습니다: {analysis_result.get('error', '알 수 없는 오류')}")

with tab2:
    st.header("분석 블록 선택")
    
    # 기본 정보나 파일 중 하나라도 있으면 진행
    has_basic_info = any([project_name, location, project_goals, additional_info])
    has_file = st.session_state.get('pdf_uploaded', False)
    
    if not has_basic_info and not has_file:
        st.warning("프로젝트 기본 정보를 입력하거나 파일을 업로드해주세요.")
        st.stop()
    
    # 예시 블록들 표시
    st.subheader("예시 분석 블록")
    
    example_blocks = get_example_blocks()
    for block in example_blocks:
        block_id = block['id']
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**{block['name']}**")
            st.caption(block['description'])
        with col2:
            if st.checkbox("선택", key=f"select_{block_id}"):
                if block_id not in st.session_state['selected_blocks']:
                    st.session_state['selected_blocks'].append(block_id)
            else:
                if block_id in st.session_state['selected_blocks']:
                    st.session_state['selected_blocks'].remove(block_id)
    
    # 사용자 정의 블록들 표시
    st.subheader("사용자 정의 블록")
    custom_blocks = load_custom_blocks()
    
    if custom_blocks:
        for block in custom_blocks:
            block_id = block['id']
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{block['name']}**")
                st.caption(block['description'])
            with col2:
                if st.checkbox("선택", key=f"select_{block_id}"):
                    if block_id not in st.session_state['selected_blocks']:
                        st.session_state['selected_blocks'].append(block_id)
                else:
                    if block_id in st.session_state['selected_blocks']:
                        st.session_state['selected_blocks'].remove(block_id)
    else:
        st.info("사용자 정의 블록이 없습니다.")
    
    # 선택된 블록들 표시 및 순서 조정
    selected_blocks = st.session_state['selected_blocks']
    if selected_blocks:
        st.success(f"{len(selected_blocks)}개 블록이 선택되었습니다:")
        
        # 선택된 블록들의 정보를 DataFrame으로 구성
        import pandas as pd
        
        block_info_list = []
        for block_id in selected_blocks:
            block_name = "알 수 없음"
            block_description = ""
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    block_description = block['description']
                    break
            block_info_list.append({
                '순서': len(block_info_list) + 1,
                '블록명': block_name,
                '설명': block_description,
                '블록ID': block_id
            })
        
        # 순서 조정을 위한 데이터프레임 생성
        df = pd.DataFrame(block_info_list)
        
        st.subheader("선택된 블록 목록 및 순서 조정")
        
        # 순서 조정 UI
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("**현재 선택된 블록들:**")
            
            # 수정 가능한 데이터 에디터로 순서 조정
            edited_df = st.data_editor(
                df[['순서', '블록명', '설명']],
                use_container_width=True,
                num_rows="fixed",
                column_config={
                    "순서": st.column_config.NumberColumn(
                        "순서",
                        help="분석 실행 순서 (숫자가 작을수록 먼저 실행)",
                        min_value=1,
                        max_value=len(block_info_list),
                        step=1
                    ),
                    "블록명": st.column_config.TextColumn(
                        "블록명",
                        disabled=True
                    ),
                    "설명": st.column_config.TextColumn(
                        "설명", 
                        disabled=True
                    )
                }
            )
        
        with col2:
            st.markdown("**빠른 순서 조정:**")
            
            # 위/아래 이동 버튼들
            for i, (_, row) in enumerate(df.iterrows()):
                st.markdown(f"**{row['블록명']}**")
                col_up, col_down = st.columns(2)
                
                with col_up:
                    if st.button("위로", key=f"up_{row['블록ID']}", disabled=(i == 0)):
                        # 위로 이동
                        if i > 0:
                            # session_state에서 직접 수정
                            current_blocks = st.session_state['selected_blocks']
                            current_blocks[i], current_blocks[i-1] = current_blocks[i-1], current_blocks[i]
                            st.session_state['selected_blocks'] = current_blocks
                            st.rerun()
                
                with col_down:
                    if st.button("아래로", key=f"down_{row['블록ID']}", disabled=(i == len(selected_blocks)-1)):
                        # 아래로 이동
                        if i < len(selected_blocks) - 1:
                            # session_state에서 직접 수정
                            current_blocks = st.session_state['selected_blocks']
                            current_blocks[i], current_blocks[i+1] = current_blocks[i+1], current_blocks[i]
                            st.session_state['selected_blocks'] = current_blocks
                            st.rerun()
        
        # 순서 변경사항이 있는지 확인하고 적용
        if not edited_df['순서'].equals(df['순서']):
            # 새로운 순서로 블록 재정렬
            new_order = edited_df.sort_values('순서')['순서'].tolist()
            
            # 블록ID와 순서를 매핑
            block_id_to_order = {}
            for i, (_, row) in enumerate(df.iterrows()):
                block_id_to_order[row['블록ID']] = new_order[i] - 1  # 0-based index
            
            # 새로운 순서로 블록들 재정렬하고 session_state에 저장
            new_blocks = sorted(st.session_state['selected_blocks'], key=lambda x: block_id_to_order[x])
            st.session_state['selected_blocks'] = new_blocks
            st.success("블록 순서가 업데이트되었습니다!")
            st.rerun()
        
        # 최종 선택된 블록들 표시
        st.subheader("최종 분석 순서")
        for i, block_id in enumerate(st.session_state['selected_blocks']):
            block_name = "알 수 없음"
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
            st.write(f"{i+1}. {block_name}")
    else:
        st.warning("분석할 블록을 선택해주세요.")

with tab3:
    st.header("분석 실행")
    
    # 기본 정보와 파일 업로드 상태 확인
    has_basic_info = any([project_name, location, project_goals, additional_info])
    has_file = st.session_state.get('pdf_uploaded', False)
    
    if not has_basic_info and not has_file:
        st.warning("프로젝트 기본 정보를 입력하거나 파일을 업로드해주세요.")
        st.stop()
    
    if not st.session_state.get('selected_blocks'):
        st.warning("먼저 분석 블록을 선택해주세요.")
        st.stop()
    
    # 입력된 정보 요약 표시
    st.subheader("분석 대상 정보")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**프로젝트 정보**")
        if project_name:
            st.write(f"• 프로젝트명: {project_name}")
        if location:
            st.write(f"• 위치/지역: {location}")
        if project_goals:
            st.write(f"• 프로젝트 목표: {project_goals[:100]}...")
        if additional_info:
            st.write(f"• 추가 정보: {additional_info[:100]}...")
    
    with col2:
        st.markdown("**파일 정보**")
        if has_file:
            file_analysis = st.session_state.get('file_analysis', {})
            # 파일명 가져오기 (session_state에서 직접 또는 uploaded_file 변수에서)
            file_name = "N/A"
            if 'uploaded_file' in st.session_state and st.session_state['uploaded_file']:
                file_name = st.session_state['uploaded_file'].name
            elif uploaded_file is not None:
                file_name = uploaded_file.name
            
            st.write(f"• 파일명: {file_name}")
            st.write(f"• 파일 유형: {file_analysis.get('file_type', 'N/A')}")
            st.write(f"• 텍스트 길이: {file_analysis.get('char_count', 0)}자")
            st.write(f"• 단어 수: {file_analysis.get('word_count', 0)}단어")
        else:
            st.write("• 파일 없음 (기본 정보만 사용)")
    
    st.markdown("---")
    
    # 블록 간 Chain of Thought 분석 (기본 활성화)
    st.info("🔗 블록 간 Chain of Thought (CoT) 분석이 활성화되어 있습니다. 각 블록의 분석 결과가 다음 블록 분석에 누적되어 연결됩니다.")
    
    if st.button("분석 시작", type="primary"):
        # DSPy 분석기 초기화
        try:
            analyzer = EnhancedArchAnalyzer()
        except Exception as e:
            st.error(f"분석기 초기화 실패: {e}")
            st.stop()
        
        # 진행 상황 표시
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        selected_blocks = st.session_state['selected_blocks']
        total_blocks = len(selected_blocks)
        
        # 파일 텍스트 가져오기
        file_text = st.session_state.get('pdf_text', '') if has_file else ""
        
        # 프로젝트 정보 구성
        project_info = {
            "project_name": project_name,
            "location": location,
            "project_goals": project_goals,
            "additional_info": additional_info,
            "file_text": file_text
        }
        
        # 블록 정보 수집
        example_blocks = get_example_blocks()
        custom_blocks = load_custom_blocks()
        block_infos = {}
        
        for block in example_blocks + custom_blocks:
            if block['id'] in selected_blocks:
                block_infos[block['id']] = block
        
        # 🔗 블록 간 Chain of Thought 분석 (항상 활성화)
        status_text.text("🔗 블록 간 Chain of Thought 분석 시작...")
        
        # 진행 상황 표시를 위한 컨테이너 생성
        progress_container = st.container()
        with progress_container:
            progress_text = st.empty()
            progress_bar = st.progress(0)
        
        # 진행 상황 콜백 함수
        def update_progress(message):
            progress_text.text(message)
            # 메시지에서 현재 블록 번호 추출하여 진행률 계산
            if "블록 분석 중" in message:
                try:
                    # "📊 1/3 블록 분석 중" 형태에서 현재 번호 추출
                    current = int(message.split("📊 ")[1].split("/")[0])
                    total = int(message.split("/")[1].split(" ")[0])
                    progress = current / total
                    progress_bar.progress(progress)
                except:
                    pass
        
        try:
            result = analyzer.analyze_blocks_with_cot(
                selected_blocks, 
                project_info, 
                file_text, 
                block_infos,
                progress_callback=update_progress
            )
            
            if result['success']:
                analysis_results = result['analysis_results']
                cot_history = result['cot_history']
                
                # 진행 상황 표시 정리
                with progress_container:
                    progress_text.text("✅ 분석 완료!")
                    progress_bar.progress(1.0)
                
                # 세션 상태에 저장
                st.session_state['analysis_results'] = analysis_results
                st.session_state['cot_history'] = cot_history
                
                # 분석 완료 메시지
                st.success(f"✅ CoT 분석 완료! {len(analysis_results)}개 블록이 연결되어 분석됨")
                
            else:
                # 진행 상황 표시 정리
                with progress_container:
                    progress_text.text("❌ 분석 실패")
                    progress_bar.progress(0)
                st.error(f"CoT 분석 실패: {result.get('error', '알 수 없는 오류')}")
                
        except Exception as e:
            # 진행 상황 표시 정리
            with progress_container:
                progress_text.text("❌ 오류 발생")
                progress_bar.progress(0)
            st.error(f"CoT 분석 중 오류 발생: {e}")
            st.error("분석을 다시 시도해주세요.")
        
        # 기본 정보와 분석 결과를 blocks.json에 저장
        import json
        from datetime import datetime
        
        # 프로젝트 정보 구성
        project_info = {
            "project_name": project_name or "프로젝트",
            "location": location or "N/A",
            "project_goals": project_goals or "N/A",
            "additional_info": additional_info or "N/A"
        }
        
        # 분석 결과를 별도 파일에 저장
        analysis_folder = "analysis_results"
        
        # analysis_results 폴더가 없으면 생성
        if not os.path.exists(analysis_folder):
            os.makedirs(analysis_folder)
        
        analysis_filename = f"analysis_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        analysis_filepath = os.path.join(analysis_folder, analysis_filename)
        
        analysis_data = {
            "project_info": project_info,
            "analysis_results": analysis_results,
            "pdf_text": st.session_state.get('pdf_text', ''),
            "file_analysis": st.session_state.get('file_analysis', {}),
            "analysis_timestamp": datetime.now().isoformat(),
            "cot_history": []  # Chain of Thought 히스토리 (향후 확장 가능)
        }
        
        # 분석 결과 파일에 저장
        try:
            with open(analysis_filepath, 'w', encoding='utf-8') as f:
                json.dump(analysis_data, f, ensure_ascii=False, indent=2)
            st.success(f"분석 결과가 {analysis_filepath}에 저장되었습니다!")
        except Exception as e:
            st.warning(f"분석 결과 저장 실패: {e}")
        
        # 결과 미리보기
        if analysis_results:
            st.subheader("분석 결과 미리보기")
            for block_id, result in analysis_results.items():
                # 블록 이름 찾기
                block_name = "알 수 없음"
                for block in example_blocks + custom_blocks:
                    if block['id'] == block_id:
                        block_name = block['name']
                        break
                
                with st.expander(f"{block_name}"):
                    st.markdown(result)

with tab4:
    st.header("결과 다운로드")
    
    if not st.session_state.get('analysis_results'):
        st.warning("먼저 분석을 실행해주세요.")
        st.stop()
    
    analysis_results = st.session_state['analysis_results']
    
    if analysis_results:
        st.success(f"{len(analysis_results)}개 분석 결과가 준비되었습니다.")
        
        # Word 문서 생성
        if st.button("Word 문서 생성", type="primary"):
            with st.spinner("Word 문서 생성 중..."):
                doc = create_word_document(project_name, analysis_results)
                
                # 메모리에 직접 바이트 데이터 생성
                import io
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_data = doc_buffer.getvalue()
                
                # 다운로드 버튼 표시
                st.download_button(
                    label="📥 Word 문서 다운로드",
                    data=file_data,
                    file_name=f"{project_name}_분석보고서.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # 개별 결과 다운로드
        st.subheader("개별 분석 결과")
        for block_id, result in analysis_results.items():
            # 블록 이름 찾기
            block_name = "알 수 없음"
            example_blocks = get_example_blocks()
            custom_blocks = load_custom_blocks()
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
            
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{block_name}**")
            with col2:
                st.download_button(
                    label="📥 다운로드",
                    data=result,
                    file_name=f"{block_name}.txt",
                    mime="text/plain",
                    key=f"download_{block_id}"
                )
    else:
        st.info("분석 결과가 없습니다.")
