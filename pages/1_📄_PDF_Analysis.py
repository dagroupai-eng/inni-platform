import streamlit as st
import os
from dotenv import load_dotenv
from pdf_analyzer import extract_text_from_pdf, analyze_pdf_content
from file_analyzer import UniversalFileAnalyzer
from dspy_analyzer import EnhancedArchAnalyzer
from prompt_processor import load_blocks, load_custom_blocks, process_prompt, save_custom_block, get_block_by_id
from docx import Document
import tempfile

# 환경변수 로드 (안전하게 처리)
try:
    load_dotenv()
except UnicodeDecodeError:
    # .env 파일에 인코딩 문제가 있는 경우 무시
    pass

# 페이지 설정
st.set_page_config(
    page_title="파일 분석",
    page_icon="📄",
    layout="wide"
)

# 제목
st.title("📄 파일 분석")
st.markdown("**건축 프로젝트 문서 분석 (PDF, Excel, CSV, 텍스트, JSON 지원)**")

# Session state 초기화
if 'project_name' not in st.session_state:
    st.session_state.project_name = ""
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'selected_blocks' not in st.session_state:
    st.session_state.selected_blocks = []
if 'pdf_text' not in st.session_state:
    st.session_state.pdf_text = ""
if 'pdf_uploaded' not in st.session_state:
    st.session_state.pdf_uploaded = False

# 블록들을 JSON 파일에서 로드
def get_example_blocks():
    """blocks.json에서 예시 블록들을 로드합니다."""
    return load_blocks()

def create_word_document(project_name, analysis_results):
    """분석 결과를 Word 문서로 생성합니다."""
    doc = Document()
    
    # 제목
    doc.add_heading(f'건축 프로젝트 분석 보고서: {project_name}', 0)
    
    # 각 분석 결과 추가
    for block_id, result in analysis_results.items():
        # 블록 이름 찾기
        block_name = "사용자 정의 블록"
        if block_id.startswith('custom_'):
            custom_blocks = load_custom_blocks()
            for block in custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        else:
            example_blocks = get_example_blocks()
            for block in example_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        
        # 섹션 제목
        doc.add_heading(block_name, level=1)
        
        # 분석 결과 내용
        doc.add_paragraph(result)
        doc.add_paragraph()  # 빈 줄
    
    return doc

# 사이드바 - 프로젝트 정보
with st.sidebar:
    st.header("📝 프로젝트 정보")
    project_name = st.text_input("프로젝트명", placeholder="예: 학생 기숙사 프로젝트", key="project_name")
    
    st.header("🔧 설정")
    
    # Streamlit secrets와 환경변수 모두 확인
    from dotenv import load_dotenv
    load_dotenv()
    
    # Streamlit secrets에서 먼저 확인
    api_key = st.secrets.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY")
    
    if not api_key:
        st.error("⚠️ ANTHROPIC_API_KEY가 설정되지 않았습니다!")
        st.info("다음 중 하나의 방법으로 API 키를 설정해주세요:")
        st.code("""
# 방법 1: .streamlit/secrets.toml 파일에 추가
[secrets]
ANTHROPIC_API_KEY = "your_api_key_here"

# 방법 2: .env 파일에 추가
ANTHROPIC_API_KEY=your_api_key_here
        """, language="toml")
        st.stop()
    else:
        st.success("✅ API 키가 설정되었습니다!")
        st.info(f"API 키 길이: {len(api_key)}자")
        st.info(f"키 소스: {'Streamlit Secrets' if st.secrets.get('ANTHROPIC_API_KEY') else '환경변수'}")

# 메인 컨텐츠
tab1, tab2, tab3, tab4 = st.tabs(["📄 파일 업로드", "🧩 분석 블록 선택", "⚡ 분석 실행", "📊 결과 다운로드"])

with tab1:
    st.header("📄 파일 업로드")
    
    uploaded_file = st.file_uploader(
        "파일을 업로드하세요",
        type=['pdf', 'xlsx', 'xls', 'csv', 'txt', 'json'],
        help="건축 프로젝트 관련 문서를 업로드하세요 (PDF, Excel, CSV, 텍스트, JSON 지원)"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ 파일 업로드 완료: {uploaded_file.name}")
        
        # 파일 확장자 확인
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        # 범용 파일 분석기 초기화
        file_analyzer = UniversalFileAnalyzer()
        
        # 파일 분석
        with st.spinner(f"{file_extension.upper()} 파일 분석 중..."):
            analysis_result = file_analyzer.analyze_file(tmp_path, file_extension)
        
        if analysis_result['success']:
            st.success(f"✅ {file_extension.upper()} 파일 분석 완료!")
            
            # 파일 정보 표시
            file_info = file_analyzer.get_file_info(tmp_path)
            st.info(f"📊 파일 정보: {file_info['file_size_mb']}MB, {analysis_result['word_count']}단어, {analysis_result['char_count']}문자")
            
            # 파일 형식별 특별 정보 표시
            if analysis_result['file_type'] == 'excel':
                st.info(f"📋 Excel 시트: {', '.join(analysis_result['sheet_names'])} ({analysis_result['sheet_count']}개 시트)")
            elif analysis_result['file_type'] == 'csv':
                st.info(f"📊 CSV 데이터: {analysis_result['shape'][0]}행 × {analysis_result['shape'][1]}열")
            
            # 세션에 저장
            st.session_state['pdf_text'] = analysis_result['text']  # 기존 변수명 유지
            st.session_state['pdf_uploaded'] = True
            st.session_state['file_type'] = analysis_result['file_type']
            st.session_state['file_analysis'] = analysis_result
            
            # 텍스트 미리보기
            with st.expander(f"📖 {file_extension.upper()} 내용 미리보기"):
                st.text(analysis_result['preview'])
        else:
            st.error(f"❌ {file_extension.upper()} 파일 분석에 실패했습니다: {analysis_result.get('error', '알 수 없는 오류')}")
        
        # 임시 파일 삭제
        os.unlink(tmp_path)

with tab2:
    st.header("🧩 분석 블록 선택")
    
    if not st.session_state.get('pdf_uploaded', False):
        st.warning("⚠️ 먼저 PDF를 업로드해주세요.")
        st.stop()
    
    # 예시 블록들 표시
    st.subheader("📚 예시 분석 블록")
    selected_blocks = []
    
    example_blocks = get_example_blocks()
    for block in example_blocks:
        block_id = block['id']
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**{block['name']}**")
            st.caption(block['description'])
        with col2:
            if st.checkbox("선택", key=f"select_{block_id}"):
                selected_blocks.append(block_id)
    
    # 사용자 정의 블록들 표시
    st.subheader("🔧 사용자 정의 블록")
    custom_blocks = load_custom_blocks()
    
    if custom_blocks:
        for block in custom_blocks:
            block_id = block['id']
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{block['name']}**")
                st.caption(block['description'])
            with col2:
                if st.checkbox("선택", key=f"select_{block_id}"):
                    selected_blocks.append(block_id)
    else:
        st.info("사용자 정의 블록이 없습니다.")
    
    # 선택된 블록들 표시
    if selected_blocks:
        st.success(f"✅ {len(selected_blocks)}개 블록이 선택되었습니다:")
        for block_id in selected_blocks:
            # 블록 이름 찾기
            block_name = "알 수 없음"
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
            st.write(f"- {block_name}")
        
        st.session_state['selected_blocks'] = selected_blocks
    else:
        st.warning("⚠️ 분석할 블록을 선택해주세요.")

with tab3:
    st.header("⚡ 분석 실행")
    
    if not st.session_state.get('pdf_uploaded', False):
        st.warning("⚠️ 먼저 PDF를 업로드해주세요.")
        st.stop()
    
    if not st.session_state.get('selected_blocks'):
        st.warning("⚠️ 먼저 분석 블록을 선택해주세요.")
        st.stop()
    
    if st.button("🚀 분석 시작", type="primary"):
        # DSPy 분석기 초기화
        try:
            analyzer = EnhancedArchAnalyzer()
        except Exception as e:
            st.error(f"❌ 분석기 초기화 실패: {e}")
            st.stop()
        
        # 분석 결과 저장용
        analysis_results = {}
        
        # 진행 상황 표시
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        selected_blocks = st.session_state['selected_blocks']
        total_blocks = len(selected_blocks)
        
        for i, block_id in enumerate(selected_blocks):
            # 블록 정보 찾기
            block_info = None
            example_blocks = get_example_blocks()
            custom_blocks = load_custom_blocks()
            
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_info = block
                    break
            
            if not block_info:
                st.error(f"블록 {block_id}를 찾을 수 없습니다.")
                continue
            
            status_text.text(f"분석 중: {block_info['name']}")
            
            # 프롬프트에 PDF 텍스트 삽입
            prompt = process_prompt(block_info, st.session_state['pdf_text'])
            
            # DSPy + CoT 분석 실행
            if block_id.startswith('custom_'):
                # 사용자 정의 블록은 custom_module 사용
                result = analyzer.analyze_custom_block(
                    prompt, 
                    st.session_state['pdf_text']
                )
            else:
                # 예시 블록은 기본 분석 사용
                result = analyzer.analyze_project(
                    {"project_name": project_name}, 
                    st.session_state['pdf_text']
                )
            
            if result['success']:
                analysis_results[block_id] = result['analysis']
                st.success(f"✅ {block_info['name']} 완료")
            else:
                st.error(f"❌ {block_info['name']} 실패: {result.get('error', '알 수 없는 오류')}")
            
            # 진행률 업데이트
            progress_bar.progress((i + 1) / total_blocks)
        
        # 분석 완료
        status_text.text("✅ 모든 분석이 완료되었습니다!")
        progress_bar.empty()
        
        # 결과를 세션에 저장
        st.session_state['analysis_results'] = analysis_results
        
        # 결과 미리보기
        if analysis_results:
            st.subheader("📊 분석 결과 미리보기")
            for block_id, result in analysis_results.items():
                # 블록 이름 찾기
                block_name = "알 수 없음"
                for block in example_blocks + custom_blocks:
                    if block['id'] == block_id:
                        block_name = block['name']
                        break
                
                with st.expander(f"📋 {block_name}"):
                    st.markdown(result)

with tab4:
    st.header("📊 결과 다운로드")
    
    if not st.session_state.get('analysis_results'):
        st.warning("⚠️ 먼저 분석을 실행해주세요.")
        st.stop()
    
    analysis_results = st.session_state['analysis_results']
    
    if analysis_results:
        st.success(f"✅ {len(analysis_results)}개 분석 결과가 준비되었습니다.")
        
        # Word 문서 생성
        if st.button("📄 Word 문서 생성", type="primary"):
            with st.spinner("Word 문서 생성 중..."):
                doc = create_word_document(project_name, analysis_results)
                
                # 임시 파일로 저장
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    doc.save(tmp_file.name)
                    tmp_path = tmp_file.name
                
                # 파일 다운로드
                with open(tmp_path, 'rb') as f:
                    st.download_button(
                        label="📥 Word 문서 다운로드",
                        data=f.read(),
                        file_name=f"{project_name}_분석보고서.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # 임시 파일 삭제
                os.unlink(tmp_path)
        
        # 개별 결과 다운로드
        st.subheader("📋 개별 분석 결과")
        for block_id, result in analysis_results.items():
            # 블록 이름 찾기
            block_name = "알 수 없음"
            example_blocks = get_example_blocks()
            custom_blocks = load_custom_blocks()
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
            
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{block_name}**")
            with col2:
                st.download_button(
                    label="📥 다운로드",
                    data=result,
                    file_name=f"{block_name}.txt",
                    mime="text/plain",
                    key=f"download_{block_id}"
                )
    else:
        st.info("분석 결과가 없습니다.")
