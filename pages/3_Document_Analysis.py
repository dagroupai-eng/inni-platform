import streamlit as st
import os
import re
import json
import hashlib
from typing import Optional, Dict, List, Any, Tuple
from collections import Counter
from dotenv import load_dotenv
from file_analyzer import UniversalFileAnalyzer
from dspy_analyzer import EnhancedArchAnalyzer, PROVIDER_CONFIG, get_api_key, get_current_provider
from prompt_processor import load_blocks, load_custom_blocks
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# 인증 모듈 import
try:
    from auth.authentication import is_authenticated, get_current_user, check_page_access
    AUTH_AVAILABLE = True
except ImportError:
    AUTH_AVAILABLE = False
try:
    from geo_data_loader import GeoDataLoader, validate_shapefile_data
    GEO_LOADER_AVAILABLE = True
except ImportError as e:
    GEO_LOADER_AVAILABLE = False
    GeoDataLoader = None
    validate_shapefile_data = None

try:  # pragma: no cover
    import streamlit_folium as st_folium
except ImportError:  # pragma: no cover
    st_folium = None

try:
    import pandas as pd
except ImportError:  # pragma: no cover
    pd = None

# 환경변수 로드 (안전하게 처리)
try:
    load_dotenv()
except UnicodeDecodeError:
    # .env 파일에 인코딩 문제가 있는 경우 무시
    pass

# 페이지 설정
st.set_page_config(
    page_title="도시 프로젝트 분석",
    page_icon=None,
    layout="wide"
)

# 세션 초기화 (로그인 + 작업 데이터 복원)
try:
    from auth.session_init import init_page_session
    init_page_session()
except Exception as e:
    print(f"세션 초기화 오류: {e}")

# ── 단계별 저장/복원 (analysis_runs / analysis_steps) ───────────────────────────
def _load_latest_steps_into_session(user_id: int, project_id: int) -> None:
    """
    현재 프로젝트의 최신 run/steps를 읽어서 session_state에 최소 결과를 복원한다.
    - 목적: 단계별 결과 표시 + 특정 단계 재실행 가능
    - 제한: cot_session 전체를 재구성하진 않음(필요 시 실행 시 새로 초기화)
    """
    try:
        from database.analysis_steps_manager import get_latest_run, list_steps

        latest = get_latest_run(project_id)
        if not latest:
            return

        run_id = latest.get("id")
        if not run_id:
            return

        steps = list_steps(run_id)
        if not steps:
            return

        # 복원된 프로젝트가 바뀔 때만 반영
        last_loaded_key = "_analysis_steps_loaded_for_project"
        if st.session_state.get(last_loaded_key) == project_id:
            return

        st.session_state[last_loaded_key] = project_id
        st.session_state["current_analysis_run_id"] = run_id

        # plan/selected_blocks 복원
        ordered_block_ids = [s.get("block_id") for s in steps if s.get("block_id")]
        if ordered_block_ids:
            st.session_state["selected_blocks"] = ordered_block_ids
            st.session_state["cot_plan"] = ordered_block_ids

        # outputs에서 결과 복원
        restored_results = {}
        restored_citations = {}
        restored_verif = {}
        skipped = []
        for s in steps:
            bid = s.get("block_id")
            status = s.get("status")
            if status == "skipped" and bid:
                skipped.append(bid)
            outp = s.get("outputs") or {}
            if bid and isinstance(outp, dict) and outp.get("analysis"):
                restored_results[bid] = outp.get("analysis")
                if outp.get("citations"):
                    restored_citations[bid] = outp.get("citations")
                if outp.get("verifications"):
                    restored_verif[bid] = outp.get("verifications")

        if restored_results:
            st.session_state.setdefault("analysis_results", {})
            st.session_state.setdefault("cot_results", {})
            st.session_state["analysis_results"].update(restored_results)
            st.session_state["cot_results"].update(restored_results)

        if restored_citations:
            st.session_state.setdefault("cot_citations", {})
            st.session_state["cot_citations"].update(restored_citations)

        if restored_verif:
            st.session_state.setdefault("cot_verifications", {})
            st.session_state["cot_verifications"].update(restored_verif)

        if skipped:
            st.session_state["skipped_blocks"] = list(dict.fromkeys(skipped))

        # 다음 실행 인덱스 계산(완료/스킵된 prefix를 건너뜀)
        completed = set(st.session_state.get("analysis_results", {}).keys())
        skipped_set = set(st.session_state.get("skipped_blocks", []))
        next_idx = 0
        for bid in ordered_block_ids:
            if bid in completed or bid in skipped_set:
                next_idx += 1
                continue
            break
        st.session_state["cot_current_index"] = min(next_idx, len(ordered_block_ids))
    except Exception as e:
        print(f"[AnalysisSteps] 복원 실패: {e}")

# 강제 불러오기 처리 (불러오기 버튼 클릭 시)
if st.session_state.get('_force_load_session'):
    try:
        from database.db_manager import execute_query
        import json

        # 플래그 제거
        del st.session_state['_force_load_session']

        # 로그인 확인
        if 'pms_current_user' in st.session_state:
            user_id = st.session_state.pms_current_user.get('id')

            if user_id:
                print(f"[강제 불러오기] 사용자 ID: {user_id}")

                # DB에서 최근 세션 조회
                result = execute_query(
                    """
                    SELECT session_data, created_at FROM analysis_sessions
                    WHERE user_id = ?
                    ORDER BY created_at DESC
                    LIMIT 1
                    """,
                    (user_id,)
                )

                if result and result[0]:
                    raw = result[0]['session_data']
                    session_data = json.loads(raw) if isinstance(raw, str) else raw
                    saved_time = result[0]['created_at']

                    print(f"[강제 불러오기] DB에서 데이터 로드: {len(session_data)}개 키")

                    # 복원 플래그 초기화 (강제 복원)
                    if 'work_session_restored_global' in st.session_state:
                        del st.session_state['work_session_restored_global']
                    if 'work_session_restoring' in st.session_state:
                        del st.session_state['work_session_restoring']

                    # 임시로 데이터 저장 (다음 rerun 때 표시용)
                    st.session_state['_loaded_data_info'] = {
                        'project_name': session_data.get('project_name', '(없음)'),
                        'location': session_data.get('location', '(없음)'),
                        'saved_time': saved_time,
                        'count': len(session_data)
                    }

                    print(f"[강제 불러오기] 복원 플래그 초기화 완료, 자동 복원 시작")
                    st.rerun()
                else:
                    st.warning("⚠️ 저장된 세션이 없습니다.")
                    print("[강제 불러오기] DB에 저장된 세션 없음")
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[강제 불러오기 오류]:\n{error_details}")
        st.error(f"❌ 불러오기 실패: {str(e)}")

# 불러오기 완료 메시지 표시
if '_loaded_data_info' in st.session_state:
    info = st.session_state['_loaded_data_info']
    del st.session_state['_loaded_data_info']
    st.success(f"✅ 저장된 정보를 불러왔습니다! (저장 시간: {info['saved_time']})")
    with st.expander("불러온 내용 확인", expanded=True):
        st.write(f"**프로젝트명**: {info['project_name']}")
        st.write(f"**위치**: {info['location']}")
        st.write(f"**총 {info['count']}개 항목 불러옴**")

# 로그인 체크
if AUTH_AVAILABLE:
    check_page_access()

# 제목
st.title("도시 프로젝트 분석")
st.markdown("**도시 프로젝트 문서 분석 (PDF, Word, Excel, CSV, 텍스트, JSON 지원)**")

# ── 프로젝트 선택 바 ──────────────────────────────────────────────────────────
try:
    from auth.project_manager import render_project_selector
    render_project_selector()
except Exception as _pm_err:
    print(f"[ProjectManager] 렌더링 오류: {_pm_err}")

# 프로젝트 선택 후 단계별 결과 자동 복원
try:
    _u = st.session_state.get("pms_current_user") or {}
    _uid = _u.get("id")
    _pid = st.session_state.get("current_project_id")
    if _uid and _pid:
        _load_latest_steps_into_session(_uid, _pid)
except Exception as _steps_err:
    print(f"[AnalysisSteps] 자동 복원 오류: {_steps_err}")

# ── 복원 알림 ─────────────────────────────────────────────────────────────────
_restore_notice = st.session_state.pop("_restore_notice", None)
if _restore_notice:
    _pn = _restore_notice.get("project_name") or "이전 프로젝트"
    _cnt = _restore_notice.get("count", 0)
    st.success(f"✅ **{_pn}** 작업 내용을 불러왔습니다. ({_cnt}개 항목 복원)", icon="📂")

# 페이지 상단 컨트롤 (리셋 버튼)
col_title, col_reset = st.columns([5, 1])
with col_reset:
    if st.button("🗑️ 페이지 초기화", use_container_width=True, help="이 페이지의 모든 데이터를 초기화합니다"):
        # Document Analysis 페이지 관련 모든 데이터 초기화
        keys_to_reset = [
            'project_name', 'location', 'latitude', 'longitude',
            'project_goals', 'additional_info', 'pdf_text', 'pdf_uploaded',
            'uploaded_file', 'file_type', 'file_analysis',
            'selected_blocks', 'analysis_results', 'cot_results',
            'cot_session', 'cot_plan', 'cot_current_index',
            'cot_running_block', 'cot_progress_messages', 'cot_feedback_inputs',
            'skipped_blocks', 'cot_citations', 'cot_history', 'cot_analyzer',
            'preprocessed_text', 'preprocessed_summary', 'preprocessing_meta',
            'reference_documents', 'reference_combined_text', 'reference_signature',
            'block_spatial_data', 'block_spatial_selection',
            'document_summary', 'doc_rag_system'
        ]
        for key in keys_to_reset:
            if key in st.session_state:
                del st.session_state[key]

        # 초기화 플래그 설정 (rerun 후 자동 복원 방지)
        st.session_state['page_just_reset'] = True
        st.session_state['work_session_restored_global'] = True
        if 'work_session_restoring' in st.session_state:
            del st.session_state['work_session_restoring']

        # 저장된 데이터는 유지 (DB/GitHub 삭제하지 않음)
        # 세션 상태만 클리어하여 화면을 깨끗하게 함
        print(f"[초기화] {len(keys_to_reset)}개 세션 키 삭제 완료 (저장된 데이터는 유지)")

        st.success("✅ 페이지가 완전히 초기화되었습니다.")
        st.rerun()

st.markdown("---")

# 사용자 인증 상태 표시 (사이드바)
if AUTH_AVAILABLE:
    with st.sidebar:
        if is_authenticated():
            user = get_current_user()
            st.success(f"로그인: {user.get('display_name', user.get('personal_number'))}")
        else:
            st.warning("로그인이 필요합니다")
            st.info("사이드바에서 '로그인' 페이지로 이동하세요.")
        st.markdown("---")

# Session state 초기화 (자동 복원 없이 빈 값으로 초기화)
if 'project_name' not in st.session_state:
    st.session_state.project_name = ""
    print("[초기화] project_name을 빈 문자열로 초기화")
else:
    print(f"[초기화] project_name 이미 존재: '{st.session_state.project_name}'")
if 'location' not in st.session_state:
    st.session_state.location = ""
if 'latitude' not in st.session_state:
    st.session_state.latitude = ""
if 'longitude' not in st.session_state:
    st.session_state.longitude = ""
if 'project_goals' not in st.session_state:
    st.session_state.project_goals = ""
if 'additional_info' not in st.session_state:
    st.session_state.additional_info = ""
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}

if 'selected_blocks' not in st.session_state:
    st.session_state.selected_blocks = []
if 'pdf_text' not in st.session_state:
    st.session_state.pdf_text = ""
if 'pdf_uploaded' not in st.session_state:
    st.session_state.pdf_uploaded = False
# 공간 데이터 초기화 (Mapping 페이지에서 업로드된 Shapefile 저장용)
if 'geo_layers' not in st.session_state:
    st.session_state.geo_layers = {}
if 'uploaded_gdf' not in st.session_state:
    st.session_state.uploaded_gdf = None
if 'uploaded_layer_info' not in st.session_state:
    st.session_state.uploaded_layer_info = None
if 'preprocessed_summary' not in st.session_state:
    st.session_state.preprocessed_summary = ""
if 'preprocessed_text' not in st.session_state:
    st.session_state.preprocessed_text = ""
if 'preprocessing_meta' not in st.session_state:
    st.session_state.preprocessing_meta = {}
if 'preprocessing_options' not in st.session_state:
    st.session_state.preprocessing_options = {
        "clean_whitespace": True,
        "collapse_blank_lines": True,
        "limit_chars": 6000,
        "include_keywords": True,
        "keyword_count": 12,
        "include_numeric_sentences": True,
        "numeric_sentence_count": 5
    }
if 'use_preprocessed_text' not in st.session_state:
    st.session_state.use_preprocessed_text = False
if 'llm_temperature' not in st.session_state:
    st.session_state.llm_temperature = 0.2
if 'llm_max_tokens' not in st.session_state:
    st.session_state.llm_max_tokens = 16000
if 'cot_session' not in st.session_state:
    st.session_state.cot_session = None
if 'cot_plan' not in st.session_state:
    st.session_state.cot_plan = []
if 'cot_current_index' not in st.session_state:
    st.session_state.cot_current_index = 0
if 'llm_provider' not in st.session_state:
    st.session_state.llm_provider = 'gemini_25flash'
if 'cot_results' not in st.session_state:
    st.session_state.cot_results = {}
if 'cot_citations' not in st.session_state:
    st.session_state.cot_citations = {}
if 'cot_progress_messages' not in st.session_state:
    st.session_state.cot_progress_messages = []
if 'cot_history' not in st.session_state:
    st.session_state.cot_history = []
if 'cot_analyzer' not in st.session_state:
    st.session_state.cot_analyzer = None
if 'cot_running_block' not in st.session_state:
    st.session_state.cot_running_block = None
if 'cot_feedback_inputs' not in st.session_state:
    st.session_state.cot_feedback_inputs = {}
if 'reference_documents' not in st.session_state:
    st.session_state.reference_documents = []
if 'reference_combined_text' not in st.session_state:
    st.session_state.reference_combined_text = ""
if 'reference_signature' not in st.session_state:
    st.session_state.reference_signature = None
if 'document_summary' not in st.session_state:
    st.session_state.document_summary = None
if 'doc_rag_system' not in st.session_state:
    st.session_state.doc_rag_system = None
if 'cot_verifications' not in st.session_state:
    st.session_state.cot_verifications = {}
if 'urban_indicator_results' not in st.session_state:
    st.session_state.urban_indicator_results = None

DEFAULT_FIXED_PROGRAM = {
    "phase1_program_intro": "",
    "phase1_program_education": "",
    "phase1_program_sports": "",
    "phase1_program_convention": "",
    "phase1_program_wellness": "",
    "phase1_program_other": ""
}

for key, value in DEFAULT_FIXED_PROGRAM.items():
    if key not in st.session_state:
        st.session_state[key] = value

def build_fixed_program_markdown() -> str:
    return "\n".join([
        "## 고정 프로그램 사양 (삼척 스포츠아카데미)",
        "",
        st.session_state.get("phase1_program_intro", "").strip(),
        "",
        "### 교육 시설",
        st.session_state.get("phase1_program_education", "").strip(),
        "",
        "### 스포츠 지원시설",
        st.session_state.get("phase1_program_sports", "").strip(),
        "",
        "### 컨벤션 시설",
        st.session_state.get("phase1_program_convention", "").strip(),
        "",
        "### 재활/웰니스",
        st.session_state.get("phase1_program_wellness", "").strip(),
        "",
        "### 기타 시설",
        st.session_state.get("phase1_program_other", "").strip()
    ])

def save_analysis_result(block_id, analysis_result, project_info=None):
    """개별 블록 분석 결과 저장 — Supabase는 save_work_session()이 처리하므로 여기서는 no-op"""
    # analysis_results는 session_state에 저장되며,
    # save_work_session() / save_analysis_progress()가 Supabase에 저장함
    return None

def load_saved_analysis_results():
    """Supabase analysis_sessions에서 현재 로그인 사용자의 분석 결과를 로드"""
    import json

    results = {}

    if 'pms_current_user' not in st.session_state:
        return results

    try:
        from database.db_manager import execute_query

        user_id = st.session_state.pms_current_user.get('id')
        if not user_id:
            return results

        project_id = st.session_state.get('current_project_id')
        if project_id:
            rows = execute_query(
                """
                SELECT session_data FROM analysis_sessions
                WHERE user_id = ? AND project_id = ?
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (user_id, project_id)
            )
        else:
            rows = execute_query(
                """
                SELECT session_data FROM analysis_sessions
                WHERE user_id = ?
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (user_id,)
            )

        if rows and rows[0]:
            raw = rows[0]['session_data']
            session_data = json.loads(raw) if isinstance(raw, str) else raw
            results = session_data.get('analysis_results', {})

    except Exception as e:
        print(f"분석 결과 로드 오류: {e}")

    return results

def get_cot_analyzer() -> Optional[EnhancedArchAnalyzer]:
    """CoT Analyzer를 가져오거나 생성합니다. Provider 변경 시 재생성합니다."""
    try:
        current_provider = get_current_provider()
        
        # Provider가 변경되었거나 analyzer가 없거나 None이면 재생성
        last_provider = st.session_state.get('_last_analyzer_provider')
        cot_analyzer_exists = st.session_state.get('cot_analyzer') is not None
        if (last_provider != current_provider) or (not cot_analyzer_exists):
            # 기존 analyzer 제거
            if 'cot_analyzer' in st.session_state:
                del st.session_state.cot_analyzer
            # 새 analyzer 생성 (예외 처리)
            try:
                analyzer = EnhancedArchAnalyzer()
                # 초기화 오류가 있는지 확인
                if hasattr(analyzer, '_init_error'):
                    init_error = analyzer._init_error
                    st.error(f"분석기 초기화 실패: {init_error}")
                    st.info(" **해결 방법**:")
                    provider_config = PROVIDER_CONFIG.get(current_provider, {})
                    api_key_env = provider_config.get('api_key_env', '')
                    display_name = provider_config.get('display_name', current_provider)
                    
                    if current_provider == 'gemini':
                        st.info("1. Google AI Studio API 키 확인:")
                        st.code(f"   .streamlit/secrets.toml 또는 환경변수에 {api_key_env} 설정", language=None)
                        st.info("2. API 키 형식 확인: AIza...로 시작하는 문자열")
                        st.info("3. Google AI Studio에서 API 키 생성: https://aistudio.google.com/app/apikey")
                    else:
                        st.info(f"1. {display_name} API 키 확인:")
                        st.code(f"   .streamlit/secrets.toml 또는 환경변수에 {api_key_env} 설정", language=None)
                        st.info("2. API 키가 올바르게 설정되었는지 확인")
                        st.info("3. Streamlit 앱을 재시작해보세요")
                    return None
                
                st.session_state.cot_analyzer = analyzer
                st.session_state._last_analyzer_provider = current_provider
            except Exception as e:
                import traceback
                error_detail = traceback.format_exc()
                st.error(f"분석기 초기화 실패: {str(e)}")
                with st.expander("상세 오류 정보", expanded=False):
                    st.code(error_detail, language='python')
                st.info(" **해결 방법**:")
                provider_config = PROVIDER_CONFIG.get(current_provider, {})
                api_key_env = provider_config.get('api_key_env', '')
                display_name = provider_config.get('display_name', current_provider)
                
                if current_provider == 'gemini':
                    st.info("1. Google AI Studio API 키 확인:")
                    st.code(f"   .streamlit/secrets.toml 또는 환경변수에 {api_key_env} 설정", language=None)
                    st.info("2. API 키 형식 확인: AIza...로 시작하는 문자열")
                    st.info("3. Google AI Studio에서 API 키 생성: https://aistudio.google.com/app/apikey")
                else:
                    st.info(f"1. {display_name} API 키 확인:")
                    st.code(f"   .streamlit/secrets.toml 또는 환경변수에 {api_key_env} 설정", language=None)
                    st.info("2. API 키가 올바르게 설정되었는지 확인")
                    st.info("3. Streamlit 앱을 재시작해보세요")
                return None
        else:
            analyzer = st.session_state.cot_analyzer
        
        # analyzer가 None인지 확인
        if analyzer is None:
            st.error("분석기가 초기화되지 않았습니다. 페이지를 새로고침하세요.")
            return None
        
        # analyzer에 초기화 오류가 있는지 확인
        if hasattr(analyzer, '_init_error'):
            st.error(f"분석기 초기화 오류: {analyzer._init_error}")
            st.info(" 위의 오류 메시지를 확인하고 API 키 설정을 확인하세요.")
            return None
        
        return analyzer
    except Exception as e:
        st.error(f"분석기 가져오기 실패: {str(e)}")
        return None

def parse_result_into_sections(text: str) -> List[Dict[str, str]]:
    """
    분석 결과 텍스트를 섹션별로 파싱합니다.
    
    Args:
        text: 분석 결과 텍스트
        
    Returns:
        섹션 리스트 (각 섹션은 {'title': str, 'content': str} 형태)
    """
    if not text:
        return [{'title': '', 'content': text}]
    
    sections = []
    lines = text.split('\n')
    current_section = {'title': '', 'content': ''}
    
    # 섹션 헤더 패턴 (##, ###, #### 등)
    section_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    
    for line in lines:
        match = section_pattern.match(line.strip())
        if match:
            # 이전 섹션 저장
            if current_section['content'].strip():
                sections.append(current_section)
            
            # 새 섹션 시작
            level = len(match.group(1))
            title = match.group(2).strip()
            # 이모지나 특수문자 제거 (탭 이름에 사용하기 위해)
            clean_title = re.sub(r'[\[\]연동대기진행완료블록결과]', '', title).strip()
            current_section = {'title': clean_title, 'content': line + '\n'}
        else:
            current_section['content'] += line + '\n'
    
    # 마지막 섹션 저장
    if current_section['content'].strip():
        sections.append(current_section)
    
    # 섹션이 없으면 전체를 하나의 섹션으로
    if not sections:
        return [{'title': '', 'content': text}]
    
    return sections

def reset_step_analysis_state(preserve_existing_results: bool = False) -> None:
    """
    단계별 분석 세션 상태를 완전히 초기화합니다.

    Args:
        preserve_existing_results: True이면 기존 분석 결과를 유지합니다.
    """
    # 분석기 내부 상태도 완전히 초기화
    try:
        EnhancedArchAnalyzer.reset_lm()
    except Exception:
        pass

    # LiteLLM 캐시 초기화 (이전 분석 결과가 캐시에서 반환되는 것 방지)
    try:
        import litellm
        if hasattr(litellm, 'cache') and litellm.cache is not None:
            litellm.cache = None
        if hasattr(litellm, '_async_client'):
            litellm._async_client = None
    except Exception:
        pass

    # DSPy 캐시 및 상태 초기화
    try:
        import dspy
        # DSPy LM 초기화 상태 리셋
        EnhancedArchAnalyzer._lm_initialized = False
        EnhancedArchAnalyzer._last_provider = None
        # DSPy 내부 캐시 초기화 시도
        if hasattr(dspy, 'cache') and dspy.cache is not None:
            if hasattr(dspy.cache, 'clear'):
                dspy.cache.clear()
        # DSPy settings 리셋
        if hasattr(dspy, 'settings'):
            try:
                dspy.settings.configure(lm=None)
            except Exception:
                pass
    except Exception:
        pass

    # 모든 세션 상태를 완전히 초기화
    st.session_state.cot_session = None
    st.session_state.cot_plan = []
    st.session_state.cot_current_index = 0
    st.session_state.cot_results = {}
    st.session_state.cot_progress_messages = []
    st.session_state.cot_running_block = None
    st.session_state.skipped_blocks = []  # 건너뛴 블록 목록 초기화
    
    # analyzer를 완전히 삭제하여 재생성되도록 함
    st.session_state.pop('cot_analyzer', None)
    st.session_state.pop('_last_analyzer_provider', None)
    
    if not preserve_existing_results:
        # 모든 분석 결과 완전히 초기화
        st.session_state.analysis_results = {}
        st.session_state.cot_citations = {}
        st.session_state.cot_history = []
        st.session_state.cot_feedback_inputs = {}
        st.session_state.document_summary = None  # 문서 요약도 초기화
        st.session_state.cot_verifications = {}
        st.session_state.doc_rag_system = None
        st.session_state.urban_indicator_results = None

        # Phase 1 관련 개별 블록 결과 초기화 (제거된 블록들)
        st.session_state.pop('phase1_requirements_cot_history', None)
        st.session_state.pop('phase1_3_requirements_text', None)
        st.session_state.pop('phase1_3_requirements_loaded', None)
        st.session_state.pop('phase1_3_selected_site', None)
        st.session_state.pop('phase1_3_selected_site_name', None)

def reset_all_state() -> None:
    """
    모든 세션 상태를 완전히 초기화합니다.
    프로젝트 정보, 파일, 분석 결과, CoT 상태, 공간 데이터 등 모든 것을 초기화합니다.
    """
    # 프로젝트 기본 정보 초기화
    st.session_state.project_name = ""
    st.session_state.location = ""
    st.session_state.project_goals = ""
    st.session_state.additional_info = ""
    
    # 파일 관련 초기화
    st.session_state.uploaded_file = None
    st.session_state.pdf_text = ""
    st.session_state.pdf_uploaded = False
    
    # 분석 결과 초기화
    st.session_state.analysis_results = {}
    st.session_state.selected_blocks = []
    
    # CoT 관련 초기화
    st.session_state.cot_session = None
    st.session_state.cot_plan = []
    st.session_state.cot_current_index = 0
    st.session_state.cot_results = {}
    st.session_state.cot_progress_messages = []
    st.session_state.cot_analyzer = None
    st.session_state.cot_running_block = None
    st.session_state.skipped_blocks = []  # 건너뛴 블록 목록 초기화
    st.session_state.cot_history = []
    st.session_state.cot_feedback_inputs = {}
    
    # 전처리 관련 초기화
    st.session_state.preprocessed_summary = ""
    st.session_state.preprocessed_text = ""
    st.session_state.preprocessing_meta = {}
    st.session_state.use_preprocessed_text = False
    st.session_state.preprocessing_options = {
        "clean_whitespace": True,
        "collapse_blank_lines": True,
        "limit_chars": 6000,
        "include_keywords": True,
        "keyword_count": 12,
        "include_numeric_sentences": True,
        "numeric_sentence_count": 5
    }
    
    # 공간 데이터 초기화
    st.session_state.geo_layers = {}
    st.session_state.uploaded_gdf = None
    st.session_state.uploaded_layer_info = None
    
    # 참고 문서 초기화
    st.session_state.reference_documents = []
    st.session_state.reference_combined_text = ""
    st.session_state.reference_signature = None
    
    # Phase 1 관련 개별 블록 결과 초기화 (제거된 블록들)
    phase1_keys = [
        'phase1_requirements_cot_history',
        'phase1_3_requirements_text',
        'phase1_3_requirements_loaded',
        'phase1_3_selected_site',
        'phase1_3_selected_site_name',
        'phase1_felo_data',
        'phase1_candidate_geo_layers',
        'phase1_candidate_sites',
        'phase1_candidate_filtered',
        'phase1_selected_sites',
        'phase1_candidate_felo_text',
        'phase1_candidate_felo_sections',
        'phase1_parse_result',
    ]
    for key in phase1_keys:
        st.session_state.pop(key, None)
    
    # 고정 프로그램 데이터 초기화
    for key in DEFAULT_FIXED_PROGRAM.keys():
        st.session_state[key] = ""

def ensure_preprocessing_options_structure(options: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """전처리 옵션 딕셔너리를 기본값과 병합합니다."""
    defaults: Dict[str, Any] = {
        "clean_whitespace": True,
        "collapse_blank_lines": True,
        "limit_chars": 6000,
        "include_keywords": True,
        "keyword_count": 12,
        "include_numeric_sentences": True,
        "numeric_sentence_count": 5,
        "include_intro_snippet": True,
        "intro_snippet_chars": 500
    }
    merged = defaults.copy()
    if options:
        for key, value in options.items():
            if key in defaults and value is not None:
                merged[key] = value
    return merged

def preprocess_analysis_text(raw_text: str, options: Dict[str, Any]) -> Tuple[str, str, Dict[str, Any]]:
    """
    분석 입력 텍스트를 전처리하고 요약 정보를 반환합니다.

    Args:
        raw_text: 원본 텍스트
        options: 전처리 옵션

    Returns:
        (정제된 텍스트, 요약 문자열, 통계 정보 딕셔너리)
    """
    if not raw_text:
        return "", "", {}
    
    import re
    
    opts = ensure_preprocessing_options_structure(options)
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    
    if opts.get("clean_whitespace", True):
        text = re.sub(r"[ \t]+", " ", text)
    
    if opts.get("collapse_blank_lines", True):
        text = re.sub(r"\n{3,}", "\n\n", text)
    
    cleaned_text = text.strip()
    limit_chars = opts.get("limit_chars")
    if isinstance(limit_chars, int) and limit_chars > 0:
        cleaned_text = cleaned_text[:limit_chars]
    
    # 단어 및 통계 계산
    word_pattern = re.compile(r"[A-Za-z가-힣0-9]{2,}")
    original_words = word_pattern.findall(raw_text)
    processed_words = word_pattern.findall(cleaned_text)
    
    # 키워드 계산
    keyword_summary = ""
    keyword_total = 0
    if opts.get("include_keywords", True) and processed_words:
        keywords = Counter(word.lower() for word in processed_words)
        keyword_count = max(1, int(opts.get("keyword_count", 10)))
        common_keywords = keywords.most_common(keyword_count)
        if common_keywords:
            keyword_total = len(common_keywords)
            keyword_summary = ", ".join(f"{word}({count})" for word, count in common_keywords)
    
    # 주요 수치 문장 추출
    numeric_summary_lines: List[str] = []
    if opts.get("include_numeric_sentences", True):
        sentences = re.split(r"(?<=[.!?])\s+|\n+", cleaned_text)
        numeric_sentences = [s.strip() for s in sentences if any(ch.isdigit() for ch in s)]
        numeric_limit = max(1, int(opts.get("numeric_sentence_count", 5)))
        for sentence in numeric_sentences[:numeric_limit]:
            numeric_summary_lines.append(f"- {sentence}")
    
    intro_snippet = ""
    if opts.get("include_intro_snippet", True) and cleaned_text:
        intro_chars = max(100, int(opts.get("intro_snippet_chars", 500)))
        intro_snippet = cleaned_text[:intro_chars]
    
    summary_sections: List[str] = []
    if intro_snippet:
        summary_sections.append(f"**요약 스니펫:**\n{intro_snippet}...")
    if keyword_summary:
        summary_sections.append(f"**핵심 키워드 Top {keyword_total}:** {keyword_summary}")
    if numeric_summary_lines:
        summary_sections.append("**주요 수치 문장:**\n" + "\n".join(numeric_summary_lines))
    
    summary_text = "\n\n".join(summary_sections).strip()
    
    stats = {
        "original_chars": len(raw_text),
        "processed_chars": len(cleaned_text),
        "original_words": len(original_words),
        "processed_words": len(processed_words),
        "keyword_summary": keyword_summary,
        "keyword_total": keyword_total
    }
    
    return cleaned_text, summary_text, stats

def get_phase1_candidate_sites():
    return st.session_state.get('phase1_candidate_sites', [])

def _safe_numeric(value, multiplier: float = 1.0):
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value) * multiplier
    text = str(value).strip()
    if not text:
        return None
    text_clean = re.sub(r'[^0-9.\-]', '', text.replace(',', ''))
    if not text_clean:
        return None
    try:
        return float(text_clean) * multiplier
    except ValueError:
        return None

def _parse_area_to_m2(value):
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text:
        return None
    if "ha" in text.lower():
        base = _safe_numeric(text)
        return base * 10000 if base is not None else None
    if "만" in text and "평" in text:
        match = re.search(r'(\d+)', text.replace(',', ''))
        if match:
            return int(match.group(1)) * 10000 * 3.3058
    if "평" in text:
        base = _safe_numeric(text)
        return base * 3.3058 if base is not None else None
    return _safe_numeric(text)

def _normalize_candidate_entry(entry):
    if not isinstance(entry, dict):
        return None
    name = entry.get("name") or entry.get("site_name") or entry.get("candidate_name")
    if not name:
        return None
    lat = _safe_numeric(entry.get("lat") or entry.get("latitude"))
    lon = _safe_numeric(entry.get("lon") or entry.get("lng") or entry.get("longitude"))
    area_m2 = _parse_area_to_m2(entry.get("area_m2") or entry.get("site_area") or entry.get("area_sq_m"))
    slope = _safe_numeric(entry.get("slope_percent") or entry.get("slope"))
    road_distance = _safe_numeric(entry.get("road_distance_m") or entry.get("road_distance"))
    facilities = entry.get("existing_facilities") or entry.get("nearby_facilities")
    try:
        facilities = int(facilities)
    except (TypeError, ValueError):
        facilities = 0
    expansion = entry.get("expansion_potential") or entry.get("expandability") or ""
    notes = entry.get("notes") or entry.get("summary") or entry.get("comment") or ""
    land_use = entry.get("land_use") or entry.get("zoning") or ""
    candidate = {
        "name": name,
        "lat": lat,
        "lon": lon,
        "area_m2": area_m2,
        "slope_percent": slope,
        "land_use": land_use,
        "road_distance_m": road_distance,
        "existing_facilities": facilities,
        "expansion_potential": expansion,
        "notes": notes
    }
    essential_fields = ["lat", "lon", "area_m2"]
    if any(candidate[field] is None for field in essential_fields):
        return None
    return candidate

def parse_candidate_sites_from_text(raw_text: str):
    if not raw_text:
        return []
    possible_texts = []
    json_pattern = re.compile(r'```json\s*(\[\s*\{.*?\}\s*\])\s*```', re.DOTALL)
    match = json_pattern.search(raw_text)
    if match:
        possible_texts.append(match.group(1))
    bracket_match = re.search(r'(\[\s*\{.*\}\s*\])', raw_text, re.DOTALL)
    if bracket_match:
        possible_texts.append(bracket_match.group(1))
    possible_texts.append(raw_text)
    for text in possible_texts:
        try:
            data = json.loads(text)
            if isinstance(data, list):
                normalized = []
                for entry in data:
                    normalized_entry = _normalize_candidate_entry(entry)
                    if normalized_entry:
                        normalized.append(normalized_entry)
                if normalized:
                    return normalized
        except Exception:
            continue
    return []

def parse_felo_candidate_blocks(raw_text: str):
    """Felo 형식의 후보지 텍스트 블록을 파싱하여 후보지 목록으로 변환"""
    if not raw_text:
        return []

    # 방법 1: 정규식으로 "후보지 X - " 패턴 찾기 (가장 안정적)
    candidate_pattern = r'후보지\s*([A-Z가-힣]+)\s*[-–—]\s*(.+?)(?=(?:🅰️|🅱️|🅲️|🅳️|🅴️|🅵️|🅶️|🅷️|🅸️|🅹️|후보지\s*[A-Z가-힣]+\s*[-–—])|$)'
    matches = list(re.finditer(candidate_pattern, raw_text, re.DOTALL))
    
    sections = []
    if matches:
        for match in matches:
            candidate_id = match.group(1).strip()  # A, B, C, D, E 등
            content = match.group(2).strip()  # 나머지 전체 내용
            
            sections.append({
                "id": candidate_id,
                "header": f"후보지 {candidate_id}",
                "content": [line.strip() for line in content.splitlines() if line.strip()]
            })

    parsed_candidates = []
    for section in sections:
        candidate_id = section.get("id", "")
        header = section["header"]
        content_list = section["content"]
        content_text = "\n".join(content_list)

        def extract_numeric_value(pattern, text, average=False, unit_multiplier=1.0):
            match = re.search(pattern, text)
            if not match:
                return None
            numbers = re.findall(r'[\d.,]+', match.group(0))
            if not numbers:
                return None
            values = []
            for num in numbers:
                try:
                    values.append(float(num.replace(',', '')))
                except ValueError:
                    continue
            if not values:
                return None
            if average and len(values) >= 2:
                return sum(values) / len(values)
            return values[0] * unit_multiplier

        area_line = next((line for line in content_list if "면적" in line), "")
        slope_line = next((line for line in content_list if "경사" in line), "")
        ic_line = next((line for line in content_list if "IC" in line or "IC 거리" in line), "")
        facility_line = next((line for line in content_list if "핵심 체육시설" in line or "체육시설" in line), "")
        constraint_line = next((line for line in content_list if line.startswith("제약")), "")
        total_score_line = next((line for line in content_list if "총점" in line), "")
        summary_lines = [line for line in content_list if line.startswith("요약")]

        area_m2 = extract_numeric_value(r"면적[^0-9]*([\d,\.]+)", area_line)
        slope_percent = extract_numeric_value(r"경사[^0-9]*([\d,\.]+(?:\s*-\s*[\d,\.]+)?)", slope_line, average=True)
        road_distance_m = extract_numeric_value(r"(IC 거리|IC)[^0-9]*([\d,\.]+)", ic_line)
        facility_distance_m = extract_numeric_value(r"핵심 체육시설[^0-9]*([\d,\.]+)", facility_line)
        score_value = extract_numeric_value(r"총점[^0-9]*([\d,\.]+)", total_score_line)
        confidence_match = re.search(r"데이터\s*신뢰도[^0-9]*([\d\.]+)/([\d\.]+)", total_score_line)
        data_confidence = None
        if confidence_match:
            try:
                numerator = float(confidence_match.group(1))
                denominator = float(confidence_match.group(2))
                if denominator > 0:
                    data_confidence = numerator / denominator
            except ValueError:
                data_confidence = None

        land_use = ""
        land_use_match = re.search(r"면적[^()]*\((.*?)\)", area_line)
        if land_use_match:
            land_use = land_use_match.group(1)

        summary_text = ""
        if summary_lines:
            summary_text = " ".join(summary_lines)
        else:
            summary_text = constraint_line

        notes = []
        if constraint_line:
            notes.append(constraint_line)
        if summary_text:
            notes.append(summary_text)
        if facility_line:
            notes.append(facility_line)

        # 후보지 이름 생성 (예: "교동·성남동 일원")
        location_name = ""
        first_line = content_list[0] if content_list else ""
        if first_line and not any(key in first_line for key in ["면적", "경사", "IC", "제약", "총점"]):
            location_name = first_line
        
        display_name = f"후보지 {candidate_id}"
        if location_name:
            display_name = f"{candidate_id} - {location_name}"
        
        candidate = {
            "name": display_name,
            "title": header,
            "area_m2": area_m2,
            "slope_percent": slope_percent,
            "road_distance_m": road_distance_m,
            "facility_distance_m": facility_distance_m,
            "existing_facilities": None,
            "expansion_potential": "보통",
            "land_use": land_use,
            "notes": "\n".join(filter(None, notes)),
            "score": score_value,
            "confidence": data_confidence,
            "summary": summary_text,
            "raw_block": content_text
        }
        parsed_candidates.append(candidate)

    return parsed_candidates

def ensure_pandas_available(feature_name: str) -> bool:
    """
    pandas 의존 기능을 사용할 수 있는지 확인.
    설치되어 있지 않으면 안내 메시지를 출력하고 False 반환.
    """
    if pd is not None:
        return True
    
    session_flag = "_pandas_warning_shown"
    if not st.session_state.get(session_flag):
        st.error(
            f"`{feature_name}` 기능을 사용하려면 pandas 라이브러리가 필요합니다. "
            "명령어 `pip install pandas` 또는 requirements 설치를 완료해주세요."
        )
        st.session_state[session_flag] = True
    return False

# 블록들을 JSON 파일에서 로드
def get_example_blocks():
    """blocks.json에서 예시 블록들을 로드합니다."""
    return load_blocks()

BLOCK_CATEGORY_MAP: Dict[str, str] = {
    "basic_info": "기본 정보 & 요구사항",
    "requirements": "기본 정보 & 요구사항",
    "project_requirements_parsing": "기본 정보 & 요구사항",
    "design_suggestions": "현황 분석 & 검증",
    "accessibility_analysis": "현황 분석 & 검증",
    "zoning_verification": "현황 분석 & 검증",
    "capacity_estimation": "현황 분석 & 검증",
    "feasibility_analysis": "사업성 & 운영 전략",
    "business_model_development": "사업성 & 운영 전략",
    "market_research_analysis": "사업성 & 운영 전략",
    "revenue_model_design": "사업성 & 운영 전략",
    "operational_efficiency_strategy": "사업성 & 운영 전략",
    "persona_scenario_analysis": "사용자 경험 & 스토리텔링",
    "storyboard_generation": "사용자 경험 & 스토리텔링",
    "customer_journey_mapping": "사용자 경험 & 스토리텔링",
}

CATEGORY_DISPLAY_ORDER: List[str] = [
    "기본 정보 & 요구사항",
    "현황 분석 & 검증",
    "사업성 & 운영 전략",
    "사용자 경험 & 스토리텔링",
    "기타",
]

def resolve_block_category(block: Dict[str, Any]) -> str:
    if not isinstance(block, dict):
        return "기타"
    category = block.get("category")
    if category:
        return category
    block_id = block.get("id")
    return BLOCK_CATEGORY_MAP.get(block_id, "기타")

def group_blocks_by_category(blocks: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for block in blocks:
        category = resolve_block_category(block)
        grouped.setdefault(category, []).append(block)
    return grouped

def iter_categories_in_order(grouped_blocks: Dict[str, List[Dict[str, Any]]]) -> List[str]:
    def _sort_key(category: str):
        if category in CATEGORY_DISPLAY_ORDER:
            return (CATEGORY_DISPLAY_ORDER.index(category), category)
        return (len(CATEGORY_DISPLAY_ORDER), category)
    return sorted(grouped_blocks.keys(), key=_sort_key)

def create_word_document(project_name, analysis_results):
    """분석 결과를 Word 문서로 생성합니다."""
    doc = Document()
    
    # 제목
    doc.add_heading(f'건축 프로젝트 분석 보고서: {project_name}', 0)
    
    # 각 분석 결과 추가
    for block_id, result in analysis_results.items():
        # 블록 이름 찾기
        block_name = "사용자 정의 블록"
        if block_id.startswith('custom_'):
            custom_blocks = load_custom_blocks()
            for block in custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        else:
            example_blocks = get_example_blocks()
            for block in example_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
        
        # 섹션 제목
        doc.add_heading(block_name, level=1)
        
        # Word 표 형식으로 처리
        add_content_with_tables(doc, result)
        doc.add_paragraph()  # 빈 줄
    
    return doc

def add_content_with_tables(doc, text):
    """텍스트를 분석하여 표는 Word 표로, 일반 텍스트는 문단으로 추가합니다."""
    import re

    # dict인 경우 (Structured Output) 문자열로 변환
    if isinstance(text, dict):
        parts = []
        if 'summary' in text:
            parts.append(str(text['summary']))
        for section in text.get('sections', []):
            if isinstance(section, dict):
                if section.get('title'):
                    parts.append(f"## {section['title']}")
                if section.get('content'):
                    parts.append(str(section['content']))
                # Structured Output의 table 필드를 마크다운 파이프 테이블로 변환
                table_data = section.get('table')
                if table_data and isinstance(table_data, dict):
                    headers = table_data.get('headers', [])
                    rows = table_data.get('rows', [])
                    if headers and rows:
                        md_lines = ['| ' + ' | '.join(str(h) for h in headers) + ' |']
                        md_lines.append('| ' + ' | '.join('---' for _ in headers) + ' |')
                        for row in rows:
                            md_lines.append('| ' + ' | '.join(str(c) for c in row) + ' |')
                        parts.append('\n'.join(md_lines))
                if section.get('table_explanation'):
                    parts.append(str(section['table_explanation']))
        if text.get('conclusion'):
            parts.append(str(text['conclusion']))
        text = '\n\n'.join(parts) if parts else str(text)
    elif not isinstance(text, str):
        text = str(text)

    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 표 시작 패턴 확인 (개선된 방식)
        if is_table_line(line):
            # 표 데이터 수집
            table_lines = [line]
            i += 1
            
            # 연속된 표 줄들 수집 (개선된 방식)
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Word 표 생성
            create_word_table(doc, table_lines)
            continue
        
        # 일반 텍스트 처리
        if line:
            # Markdown 헤더 처리
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                doc.add_heading(header_text, level=min(level, 6))
            else:
                # 리스트 처리
                if line.startswith('- '):
                    line = '• ' + line[2:]
                elif line.startswith('* '):
                    line = '• ' + line[2:]
                
                # 볼드 텍스트 처리 (**text**)
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                
                doc.add_paragraph(line)
        
        i += 1

def create_word_table(doc, table_lines):
    """Markdown 표 줄들을 Word 표로 변환합니다."""
    if not table_lines:
        return
    
    # 표 데이터 파싱
    table_data = []
    for line in table_lines:
        # |로 구분된 셀들 추출
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # 첫 번째와 마지막 빈 요소 제거
        if cells:
            table_data.append(cells)
    
    if not table_data:
        return
    
    # 첫 번째 행이 헤더 구분선인지 확인 (--- 형태)
    if len(table_data) > 1 and all(cell == '---' or cell == '------' or cell == '' for cell in table_data[1]):
        headers = table_data[0]
        data_rows = table_data[2:]
    else:
        headers = None
        data_rows = table_data
    
    # 열 수 결정
    max_cols = max(len(row) for row in table_data) if table_data else 2
    
    # Word 표 생성 - 개선된 방식
    try:
        table = doc.add_table(rows=len(data_rows) + (1 if headers else 0), cols=max_cols)
        table.style = 'Table Grid'
        
        # 표 자동 크기 조절 활성화
        table.allow_autofit = True
        table.autofit = True
        
        # 헤더 추가
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                if i < len(header_row.cells):
                    cell = header_row.cells[i]
                    cell.text = clean_text_for_pdf(header)
                    
                    # 헤더 스타일링 강화
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                        # 셀 패딩 조정
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(2)
        
        # 데이터 행 추가
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data_rows):
            if start_row + i < len(table.rows):
                table_row = table.rows[start_row + i]
                for j, cell_data in enumerate(row_data):
                    if j < len(table_row.cells):
                        cell = table_row.cells[j]
                        cell.text = clean_text_for_pdf(cell_data)
                        
                        # 셀 스타일링
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                            # 셀 패딩 조정
                            paragraph.paragraph_format.space_before = Pt(1)
                            paragraph.paragraph_format.space_after = Pt(1)
        
        # 표 후 빈 줄 추가
        doc.add_paragraph()
        
    except Exception as e:
        print(f"Word 표 생성 오류: {e}")
        # 오류 발생 시 텍스트로 대체
        doc.add_paragraph("[표 생성 실패 - 원본 데이터]")
        for row in table_data:
            doc.add_paragraph(" | ".join(row))
        doc.add_paragraph()

def is_table_line(line):
    """한 줄이 표 행인지 확인"""
    if not line:
        return False

    # | 구분자가 있는 경우 (마크다운 표 형식)
    if '|' in line and line.count('|') >= 2:
        return True

    # 탭으로 구분된 경우
    if '\t' in line:
        return True

    # 2개 이상의 공백으로 구분된 경우 (정렬된 텍스트)
    if re.search(r'\s{2,}', line):
        return True

    return False

def render_structured_response(response: dict):
    """JSON 구조화된 응답을 Streamlit으로 렌더링합니다.

    Args:
        response: AnalysisResponse 스키마를 따르는 딕셔너리
            - summary: 요약 텍스트
            - sections: 섹션 리스트
            - conclusion: 결론 (선택)
    """
    if not response or not isinstance(response, dict):
        st.warning("응답 데이터가 올바르지 않습니다.")
        return

    # 요약 렌더링
    summary = response.get('summary', '')
    if summary:
        st.markdown("### 📋 분석 요약")
        st.markdown(summary)
        st.markdown("---")

    # 섹션별 렌더링
    sections = response.get('sections', [])
    for idx, section in enumerate(sections):
        if not isinstance(section, dict):
            continue

        title = section.get('title', f'섹션 {idx + 1}')
        content = section.get('content', '')
        table_data = section.get('table')
        table_explanation = section.get('table_explanation', '')

        # 섹션 제목
        st.markdown(f"### {title}")

        # 섹션 본문
        if content:
            st.markdown(content)

        # 표 렌더링
        if table_data and isinstance(table_data, dict):
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])
            caption = table_data.get('caption', '')

            if headers and rows:
                try:
                    # 행의 열 개수를 헤더에 맞춤
                    max_cols = len(headers)
                    normalized_rows = []
                    for row in rows:
                        if len(row) < max_cols:
                            row = list(row) + [''] * (max_cols - len(row))
                        elif len(row) > max_cols:
                            row = row[:max_cols]
                        normalized_rows.append(row)

                    df = pd.DataFrame(normalized_rows, columns=headers)

                    if caption:
                        st.caption(caption)

                    st.dataframe(df, use_container_width=True, hide_index=True)
                except Exception as e:
                    st.error(f"표 렌더링 오류: {e}")
                    # 폴백: 원본 데이터 표시
                    st.json(table_data)

        # 표 해설
        if table_explanation:
            st.markdown(f"**[표 해설]** {table_explanation}")

        st.markdown("")  # 섹션 간 여백

    # 결론 렌더링
    conclusion = response.get('conclusion', '')
    if conclusion:
        st.markdown("---")
        st.markdown("### 🎯 결론")
        st.markdown(conclusion)


def _try_parse_structured_json(text: str):
    """문자열에서 마크다운 코드블록을 제거한 뒤 JSON 파싱을 시도하고, summary/sections를 가진 dict면 반환."""
    if not text or not isinstance(text, str):
        return None
    s = text.strip()
    # 앞뒤 마크다운 코드블록 제거 (```json ... ``` 또는 ``` ... ```)
    s = re.sub(r'^\s*```(?:json)?\s*\n?', '', s)
    s = re.sub(r'\n?\s*```\s*$', '', s)
    s = s.strip()
    if not s:
        return None
    try:
        import json
        parsed = json.loads(s)
        if isinstance(parsed, dict) and 'summary' in parsed and 'sections' in parsed:
            return parsed
    except (json.JSONDecodeError, TypeError):
        pass
    return None


def render_analysis_result(result):
    """분석 결과를 자동으로 감지하여 적절한 방식으로 렌더링합니다.

    - dict이고 'sections' 키가 있으면: render_structured_response 사용
    - 문자열이 JSON(summary/sections) 형태면 파싱 후 구조화 렌더링 (코드블록 포함)
    - 그 외: render_markdown_with_tables 사용
    """
    if isinstance(result, dict) and 'sections' in result:
        render_structured_response(result)
    elif isinstance(result, str):
        parsed = _try_parse_structured_json(result)
        if parsed is not None:
            render_structured_response(parsed)
            return
        render_markdown_with_tables(result)
    elif isinstance(result, dict) and 'analysis' in result:
        # 분석 결과가 래핑된 경우
        analysis = result.get('analysis')
        if isinstance(analysis, dict) and 'sections' in analysis:
            render_structured_response(analysis)
        else:
            render_markdown_with_tables(str(analysis) if analysis else "")
    else:
        st.warning("알 수 없는 결과 형식입니다.")
        st.json(result) if isinstance(result, dict) else st.text(str(result))


def render_markdown_with_tables(text):
    """마크다운 텍스트를 렌더링하면서 테이블은 st.dataframe()으로 변환합니다."""
    if not text or not isinstance(text, str):
        return

    lines = text.split('\n')
    i = 0
    buffer = []  # 일반 텍스트 버퍼

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 마크다운 테이블 시작 감지 (| 구분자 기준)
        if '|' in stripped and stripped.count('|') >= 2:
            # 버퍼에 있는 텍스트 먼저 출력
            if buffer:
                st.markdown('\n'.join(buffer))
                buffer = []

            # 테이블 라인 수집
            table_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if '|' in next_line and next_line.count('|') >= 2:
                    table_lines.append(next_line)
                    i += 1
                else:
                    break

            # 테이블을 DataFrame으로 변환
            if len(table_lines) >= 2 and pd is not None:
                try:
                    # 파싱 - 구분선(--- 패턴) 행은 제외
                    parsed_rows = []
                    for tl in table_lines:
                        parts = tl.split('|')
                        # 시작과 끝의 빈 문자열 제거 (| 로 시작하거나 끝나는 경우)
                        if parts and parts[0].strip() == '':
                            parts = parts[1:]
                        if parts and parts[-1].strip() == '':
                            parts = parts[:-1]
                        cells = [c.strip() for c in parts]

                        # 구분선 행인지 확인 (모든 셀이 ---, :---, ---:, :---: 패턴이거나 빈 경우)
                        if cells:
                            is_separator_row = all(
                                re.match(r'^[-:]+$', c) or c == ''
                                for c in cells
                            )
                            # 구분선 행은 건너뛰기
                            if not is_separator_row:
                                parsed_rows.append(cells)

                    if len(parsed_rows) >= 1:
                        # 첫 번째 행을 헤더로, 나머지를 데이터로 사용
                        headers = parsed_rows[0]
                        data = parsed_rows[1:] if len(parsed_rows) > 1 else []

                        # DataFrame 생성
                        if data:
                            max_cols = len(headers)
                            normalized_data = []
                            for row in data:
                                if len(row) < max_cols:
                                    row = row + [''] * (max_cols - len(row))
                                elif len(row) > max_cols:
                                    row = row[:max_cols]
                                normalized_data.append(row)

                            df = pd.DataFrame(normalized_data, columns=headers)
                            st.dataframe(df, use_container_width=True, hide_index=True)
                            continue
                        elif headers:
                            # 데이터가 없고 헤더만 있는 경우
                            df = pd.DataFrame(columns=headers)
                            st.dataframe(df, use_container_width=True, hide_index=True)
                            continue
                except Exception as e:
                    pass

            # 파싱 실패 시 원본 출력
            st.code('\n'.join(table_lines), language=None)
            continue

        # 일반 라인은 버퍼에 추가
        buffer.append(line)
        i += 1

    # 남은 버퍼 출력
    if buffer:
        st.markdown('\n'.join(buffer))

def is_table_format(text):
    """텍스트가 표 형식인지 확인"""
    try:
        if not text or not isinstance(text, str):
            return False
            
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False
        
        # 1. 마크다운 표 형식 확인 (| 구분자)
        pipe_count = text.count('|')
        if pipe_count >= 3:  # 최소 1x2 표를 위해서는 3개의 | 필요
            # 구분선이 있는지 확인 (표의 특징)
            for line in lines:
                line = line.strip()
                if re.match(r'^[\s\-=_:|]+\s*$', line):
                    return True
            # 구분선이 없어도 |가 많이 있으면 표로 간주
            if pipe_count >= 6:
                return True
        
        # 2. 구분선 확인 (마크다운 표 구분선)
        for line in lines:
            line = line.strip()
            if re.match(r'^[\s\-=_:|]+\s*$', line):
                return True
        
        # 3. 탭 구분자 확인
        tab_count = sum(1 for line in lines if '\t' in line)
        if tab_count >= 2:
            return True
        
        return False
        
    except Exception as e:
        print(f"표 형식 확인 오류: {e}")
        return False

def clean_text_for_pdf(text):
    """PDF/Word용 텍스트 정리"""
    if not text:
        return ""
    
    import re
    
    # HTML 태그 제거
    text = re.sub(r'<[^>]+>', '', text)
    
    # Markdown 볼드 제거 (**text** -> text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Markdown 이탤릭 제거 (*text* -> text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # 특수 문자 정리
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    
    # 연속된 공백 정리
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def _extract_first_number(text, pattern, default=None, transform=None):
    if not text:
        return default
    match = re.search(pattern, text, re.IGNORECASE)
    if not match:
        return default
    value_text = match.group(1)
    try:
        value = float(value_text.replace(',', '').replace(' ', ''))
    except ValueError:
        return default
    return transform(value) if transform else value

def _parse_area_requirement(structured_text):
    default_area = 330000
    if not structured_text:
        return default_area
    # ㎡ 단위 우선 탐색
    area = _extract_first_number(structured_text, r'([\d,]+)\s*㎡', default=None)
    if area:
        return int(area)
    # "만 평" 형태 탐색
    match = re.search(r'(\d+)\s*만\s*평', structured_text)
    if match:
        area_pyong = int(match.group(1)) * 10000
        return int(area_pyong * 3.3058)
    # 일반 "평" 형태 탐색
    match = re.search(r'([\d,]+)\s*평', structured_text)
    if match:
        area_pyong = float(match.group(1).replace(',', ''))
        return int(area_pyong * 3.3058)
    return default_area

def _parse_slope_requirement(structured_text):
    return _extract_first_number(structured_text, r'경사도[^0-9]*([\d.]+)\s*%', default=5.0)

def _parse_road_requirement(structured_text):
    return _extract_first_number(structured_text, r'도로[^0-9]*([\d,]+)\s*km', default=None, transform=lambda v: int(v * 1000))

def _parse_priority_weights(structured_text):
    default_weights = {
        "접근성": 30,
        "연계성": 25,
        "확장성": 20,
        "경제성": 15,
        "공공성": 10
    }
    if not structured_text:
        return default_weights
    pattern = r'([가-힣A-Za-z]+)\s*\((\d+)%\)'
    matches = re.findall(pattern, structured_text)
    found = {}
    for name, perc in matches:
        try:
            value = int(perc)
        except ValueError:
            continue
        for key in default_weights.keys():
            if key in name:
                found[key] = value
                break
    return {**default_weights, **found}

def derive_phase1_defaults():
    structured_text = st.session_state.get('phase1_requirements_structured', '')
    defaults = {
        "min_area": _parse_area_requirement(structured_text),
        "max_slope": _parse_slope_requirement(structured_text),
        "max_road_distance": _parse_road_requirement(structured_text) or 1000,
        "weights": _parse_priority_weights(structured_text)
    }
    return defaults

def filter_candidate_sites(min_area, max_slope, max_road_distance, include_expansion, sites=None):
    if sites is None:
        sites = get_phase1_candidate_sites()
    if not sites:
        return []
    filtered = []
    for site in sites:
        area_m2 = site.get("area_m2")
        slope_percent = site.get("slope_percent")
        road_distance = site.get("road_distance_m")
        if area_m2 is None or area_m2 < min_area:
            continue
        if slope_percent is None or slope_percent > max_slope:
            continue
        if road_distance is None or road_distance > max_road_distance:
            continue
        if include_expansion and site.get("expansion_potential") not in ["우수", "보통", "높음", "중간"]:
            continue
        filtered.append(site)
    return filtered

# Phase 1.3에서 AI가 생성한 시설 목록을 사용하므로 하드코딩된 FACILITY_LIBRARY는 제거됨

def render_phase1_1(project_name, location, project_goals, additional_info):
    st.markdown("### Mission 1 · Phase 1.1 — 요구사항 정리")
    st.caption("🟨 학생 입력 → 🟦 자체 프로그램\n\n1) 고정 프로그램 사양 확인\n2) 학생이 워크시트 내용을 자유롭게 입력하고\n3) 프로그램이 구조화된 요구사항 요약(블록 1)과 데이터 체크리스트(블록 2)를 생성합니다.")

    with st.expander("고정 프로그램 사양 (삼척 스포츠아카데미)", expanded=False):
        st.write("아래 항목은 학생이 직접 수정할 수 있으며, 블록 1과 블록 2 실행 시 함께 전달됩니다.")
        st.text_area(
            "도입 설명",
            key="phase1_program_intro",
            height=80,
            placeholder="예: 삼척 스포츠아카데미의 핵심 방향성과 기본 요구사항을 간단히 입력하세요.",
            help="프로그램 전반에 대한 소개나 주의사항을 입력하세요."
        )
        st.markdown("#### 교육 시설")
        st.text_area(
            "교육 시설 항목",
            key="phase1_program_education",
            height=120,
            placeholder="- 학교: 국제학교(중/고)와 국내 고등학교 (학년 당 약 100명 정원)\n- 부대시설: 식당, 생활관, 강당 등\n- 참고 사례: 제주 국제학교 및 서울체육고등학교",
            help="교육 시설 관련 요구사항을 자유롭게 입력하세요."
        )
        st.markdown("#### 스포츠 지원시설")
        st.text_area(
            "스포츠 지원시설 항목",
            key="phase1_program_sports",
            height=140,
            placeholder="- 핵심종목: 테니스, 양궁, 배드민턴, 펜싱\n- 추가종목: 아이스하키, 컬링 등\n- 확장종목: 러닝마라톤, 야구, 축구 등\n- 확장 전략: 단계적 확대 계획",
            help="핵심/추가/확장 종목 등을 입력하세요."
        )
        st.markdown("#### 컨벤션 시설")
        st.text_area(
            "컨벤션 시설 항목",
            key="phase1_program_convention",
            height=100,
            placeholder="- 200실 규모 호텔\n- 국제 컨벤션 홀\n- 선수/방문객 편의 리테일 시설",
            help="호텔, 컨벤션, 리테일 등 방문객 관련 시설을 입력하세요."
        )
        st.markdown("#### 재활/웰니스")
        st.text_area(
            "재활/웰니스 항목",
            key="phase1_program_wellness",
            height=100,
            placeholder="- 스포츠 의학·재활센터\n- 웰니스 프로그램 및 기업 입주시설",
            help="재활센터, 웰니스 프로그램 등의 요구사항을 입력하세요."
        )
        st.markdown("#### 기타 시설")
        st.text_area(
            "기타 시설 항목",
            key="phase1_program_other",
            height=120,
            placeholder="- 주민 개방시설\n- 국제 캠프장, 축제 관련 시설 등",
            help="그 밖에 포함하고 싶은 기타 시설을 입력하세요."
        )

    with st.expander("단계 1-1-1 · 프로젝트 요구사항 워크시트 입력", expanded=not st.session_state.get('phase1_requirements_structured')):
        st.markdown(
            """**워크시트 입력 템플릿**

    **1. 프로젝트 목표 (필수)**  
    우리가 만들고 싶은 아카데미는? (최소 1가지 이상 작성)  
    □ 엘리트 선수를 키워서 해외 진출시키는 곳  
    □ 지역 주민도 함께 즐길 수 있는 개방된 공간  
    □ 사계절 운영 가능한 복합시설  
    □ 기타: _________________________________  
    □ 기타: _________________________________  
    → 우리 팀 핵심: "_______________________"  

    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━  

    **2. 규모**  
    - 학생 수: 약 _____명 (필수)  
    TIP: 100-300명 사이가 일반적  
    - 면적: 대략 _____만 평 (선택)  
    TIP: 모르면 비워두세요  
    - 예산: _____억 (선택)  
    TIP: "모르겠음" 또는 "제약 없음" 선택 가능  

    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━  

    **3. 제약조건 (선택)**  
    우리 프로젝트에서 꼭 고려해야 할 제약은?  
    아래 예시 참고해서 자유롭게 작성:  
    - 삼척시의 인구가 적어서 지역만으로는 학생 모집이 어려움  
    - 겨울에 관광객이 적어서 수익이 떨어질 수 있음  
    - 산지가 많아서 평평한 땅이 부족함  
    우리 팀 제약:  
    □ _________________________________  
    □ _________________________________  
    □ _________________________________  

    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━  

    **4. 우선순위 (필수 - 최소 3개 선택)**  
    무엇이 가장 중요한가? 번호를 매겨보세요.  
    [  ] 교통 접근성 (고속도로, 역, 공항)  
    [  ] 기존 체육시설과의 연계  
    [  ] 확장 가능성  
    [  ] 경제성 (저렴한 토지)  
    [  ] 주민 접근성  
    [  ] 좋은 경관/환경  
    [  ] 기타: _____________  
    TIP:  
    - 1, 2, 3... 순서대로 번호 매기기  
    - 최소 3개는 선택해주세요  
    - "왜 이게 중요한가?" 간단히 메모해도 좋아요  

    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━  

    **5. 핵심 가치 (필수)**  
    이 아카데미를 한 문장으로 표현하면?  
    예시:  
    - "세계 무대로 가는 지름길"  
    - "지역과 함께 성장하는 스포츠 허브"  
    - "데이터로 만드는 차세대 스포츠 교육"  
    → _________________________________________"""
        )
        st.text_area(
            "워크시트 전체 내용을 입력하세요",
            key="phase1_requirements_text",
            placeholder="예) 수용 인원, 목표 면적, 운영 목표, 제약조건 등을 자유롭게 정리해주세요.",
            height=220,
            help="학생들이 정리한 요구사항 워크시트 전체 내용을 붙여넣으세요."
        )
        col_input_actions = st.columns([1, 1])
        with col_input_actions[1]:
            if st.button("입력 초기화", key="reset_phase1_requirements"):
                st.session_state['phase1_requirements_text'] = ""
                st.session_state.pop('phase1_requirements_cot_history', None)
                st.rerun()

    # 제거된 블록들: phase1_requirements_structuring, phase1_data_inventory
    # 관련 섹션 제거됨

    if False:  # 제거된 블록 조건
        with st.expander("📤 Felo AI 전달 데이터 정리", expanded=False):
            st.markdown("""
            **이 데이터는 Felo AI로 전달하여 후보지를 추출하는 데 사용됩니다.**
            
            아래 내용을 복사하여 Felo AI에 전달하세요.
            """)

            fixed_program = "\n".join([
                build_fixed_program_markdown(),
                "",
                "## 입지 선정 목표",
                "삼척 스포츠아카데미의 최적 입지와 규모를 선정하고, 관련 프로그램을 확정",
                "",
                "## 활용 데이터",
                "- 입지 특징",
                "- GIS 공간 데이터",
                "- 교육발전특구 법규",
                "- 접근성 분석",
                "- 도시재생 잠재력 데이터",
                "",
                "## 최종 목표",
                "시각화된 데이터 기반의 입지 선정 근거 마련"
            ])

            block1_result = st.session_state.get('phase1_requirements_structured', '')
            block2_result = st.session_state.get('phase1_data_inventory', '')

            felo_data = f"""{fixed_program}

        ---

        ## 요구사항 구조화 결과 (블록 1)

        {block1_result}

        ---

        ## Felo AI 분석 요청사항

        위 요구사항과 데이터 목록을 바탕으로 삼척시 내 최적 후보지를 추출해주세요.

        ### 분석 조건
        - 입지 특징 분석
        - GIS 공간 데이터 활용
        - 교육발전특구 법규 검토
        - 접근성 평가
        - 도시재생 잠재력 평가

        ### 출력 형식
        - 후보지 목록 (위치, 면적, 경사도, 도로 접근성 등)
        - 각 후보지별 평가 점수
        - 시각화된 지도 데이터
        """

            st.session_state['phase1_felo_data'] = felo_data

            st.text_area(
                "Felo AI 전달 데이터 (복사하여 사용)",
                value=felo_data,
                height=400,
                key="felo_data_display",
                help="전체 내용을 복사하여 Felo AI에 전달하세요."
            )

            st.download_button(
                label="Felo AI 전달 데이터 다운로드",
                data=felo_data,
                file_name="felo_ai_input_data.txt",
                mime="text/plain",
                key="download_felo_data"
            )

            st.info(" 이 데이터를 Felo AI에 전달하면 후보지 추출 결과를 받을 수 있습니다.")

def render_phase1_2(project_name, location, project_goals, additional_info):
    st.markdown("### Mission 1 · Phase 1.2 — 후보지 탐색 & 검토")

    # 1. 후보지 공간 데이터 (Shapefile 업로드)
    with st.expander("후보지 공간 데이터 (Shapefile 업로드)", expanded=False):
        st.caption("Felo 또는 외부 분석에서 받은 후보지 Shapefile(ZIP)을 업로드하면 지도에서 시각화할 수 있습니다.")
        uploaded_shapefiles = st.file_uploader(
            "Shapefile ZIP 업로드 (복수 선택 가능)",
            type=["zip"],
            accept_multiple_files=True,
            key="phase1_candidate_shapefiles"
        )

        if uploaded_shapefiles:
            if not GEO_LOADER_AVAILABLE or GeoDataLoader is None:
                st.error("GeoDataLoader를 사용할 수 없습니다. geopandas가 설치되지 않았습니다.")
                st.info("""
                **설치 방법:**
                
                1. **conda 사용 (권장):**
                   ```
                   conda install -c conda-forge geopandas
                   ```
                
                2. **pip 사용:**
                   ```
                   pip install geopandas shapely pyproj
                   ```
                
                3. **install.bat 실행:**
                   프로젝트 루트에서 `install.bat`을 실행하면 자동으로 설치됩니다.
                """)
            else:
                loader = GeoDataLoader()
                loaded = 0
                errors = []
                with st.spinner(f"{len(uploaded_shapefiles)}개 파일 처리 중..."):
                    for uploaded in uploaded_shapefiles:
                        layer_name = uploaded.name.replace(".zip", "").replace(".ZIP", "")
                        result = loader.load_shapefile_from_zip(uploaded.getvalue(), encoding="cp949")
                        if result.get("success"):
                            validation = validate_shapefile_data(result["gdf"])
                        if validation.get("valid", False):
                            st.session_state['phase1_candidate_geo_layers'][layer_name] = {
                                "gdf": result["gdf"],
                                "info": result
                            }
                            loaded += 1
                        else:
                            issues = ", ".join(validation.get("issues", []))
                            errors.append(f" {layer_name}: {issues or '데이터 검증 실패'}")
                    else:
                        errors.append(f"[실패] {layer_name}: {result.get('error', '알 수 없는 오류')}")
            if loaded:
                st.success(f" {loaded}개 레이어를 불러왔습니다.")
            for err in errors:
                st.warning(err)

        if st.session_state.get('phase1_candidate_geo_layers'):
            st.markdown("##### 업로드된 레이어")
            for layer_name, layer_data in list(st.session_state['phase1_candidate_geo_layers'].items()):
                with st.expander(f"📂 {layer_name}", expanded=False):
                    info = layer_data.get("info", {})
                    st.write(f"- 피처 수: {info.get('feature_count', 0):,}개")
                    st.write(f"- 좌표계: {info.get('crs', 'Unknown')}")
                    st.write(f"- 컬럼 수: {len(info.get('columns', []))}개")
                    if st.button("레이어 삭제", key=f"phase1_candidate_layer_delete_{layer_name}"):
                        del st.session_state['phase1_candidate_geo_layers'][layer_name]
                        st.rerun()

            if st.session_state.get('phase1_candidate_geo_layers'):
                if GEO_LOADER_AVAILABLE and GeoDataLoader is not None:
                    loader = GeoDataLoader()
                    st.markdown("##### 지도 시각화")
                    geo_layers_dict = {
                        lname: ldata["gdf"]
                        for lname, ldata in st.session_state['phase1_candidate_geo_layers'].items()
                    }
                    folium_map = None
                    if st_folium:
                        try:
                            folium_map = loader.create_folium_map_multilayer(geo_layers_dict)
                        except Exception as map_error:  # pragma: no cover
                            st.warning(f"Folium 지도를 생성하는 중 오류가 발생했습니다: {map_error}")
                            folium_map = None
                    if folium_map and st_folium:
                        st_folium.st_folium(folium_map, width=1100, height=540)
                    else:
                        if pd is None:
                            st.warning("pandas가 설치되어 있지 않아 간단 지도를 표시할 수 없습니다.")
                            st.info("`pip install pandas`를 실행한 뒤 다시 시도해주세요.")
                        else:
                            fallback_frames = []
                            for gdf in geo_layers_dict.values():
                                df_for_map = loader.gdf_to_dataframe_for_map(gdf)
                                if not df_for_map.empty:
                                    fallback_frames.append(df_for_map.head(500))
                            if fallback_frames:
                                fallback_df = pd.concat(fallback_frames, ignore_index=True)
                                st.map(fallback_df, size=20)
                            else:
                                st.info("표시 가능한 좌표 데이터가 없습니다. `pip install streamlit-folium folium` 설치 후 다시 시도하세요.")

    # 2. 🗒️ Felo 후보지 텍스트 붙여넣기
    with st.expander("🗒️ Felo 후보지 텍스트 붙여넣기", expanded=False):
        st.caption("Felo AI에서 받은 후보지 요약을 그대로 붙여넣으면 자동으로 표준화됩니다.")
        st.session_state['phase1_candidate_felo_text'] = st.text_area(
            "Felo 후보지 텍스트",
            value=st.session_state.get('phase1_candidate_felo_text', ''),
            height=260,
            placeholder="예) 🅰️ 후보지 A - 교동·성남동 일원 ...",
            key="phase1_candidate_felo_text_area"
        )
        col_felo_actions = st.columns([1, 1])
        with col_felo_actions[0]:
            if st.button("텍스트 파싱", key="phase1_parse_felo_blocks"):
                felo_text = st.session_state.get('phase1_candidate_felo_text', '')
                parsed = parse_felo_candidate_blocks(felo_text)
                if parsed:
                    st.session_state['phase1_candidate_sites'] = parsed
                    st.session_state['phase1_candidate_filtered'] = None
                    st.session_state['phase1_selected_sites'] = [entry["name"] for entry in parsed]
                    st.session_state['phase1_candidate_felo_sections'] = parsed
                    st.session_state['phase1_parse_result'] = f"{len(parsed)}개 파싱 완료"
                    st.rerun()
                else:
                    st.warning("입력된 텍스트에서 후보지를 찾을 수 없습니다. 형식을 확인해주세요.")
        
        # 파싱 결과 표시 (rerun 후에도 보이도록)
        if st.session_state.get('phase1_parse_result'):
            st.success(st.session_state['phase1_parse_result'])
            parsed_sites = st.session_state.get('phase1_candidate_sites', [])
            if parsed_sites:
                st.write("**파싱된 후보지:**")
                for idx, site in enumerate(parsed_sites, 1):
                    st.write(f"{idx}. {site.get('name', 'N/A')}")
        with col_felo_actions[1]:
            if st.button("입력 초기화", key="phase1_reset_felo_blocks"):
                st.session_state['phase1_candidate_felo_text'] = ""
                st.session_state['phase1_candidate_felo_sections'] = []
                st.session_state.pop('phase1_candidate_sites', None)
                st.session_state.pop('phase1_candidate_filtered', None)
                st.rerun()

    # 3. 📍 후보지 목록
    candidate_sites = st.session_state.get('phase1_candidate_sites', [])
    
    if not candidate_sites:
        st.info("후보지 데이터가 없습니다. Shapefile을 업로드하거나 Felo 텍스트를 붙여넣어주세요.")
        return
    
    st.markdown("#### 📍 후보지 목록")
    if pd is not None:
        df_sites = pd.DataFrame(candidate_sites)
        df_display = df_sites.rename(columns={
            "name": "후보지",
            "area_m2": "면적(㎡)",
            "slope_percent": "경사도(%)",
            "land_use": "토지용도",
            "road_distance_m": "도로거리(m)",
            "existing_facilities": "주변 체육시설(개소)",
            "expansion_potential": "확장 잠재력",
            "facility_distance_m": "핵심시설 거리(m)",
            "score": "총점(점)",
            "confidence": "데이터 신뢰도",
            "notes": "메모"
        })
        st.dataframe(df_display, use_container_width=True)
        
        # 지도 표시
        lat_series = df_sites.get("lat")
        lon_series = df_sites.get("lon")
        if lat_series is not None and lon_series is not None:
            if lat_series.notnull().any() and lon_series.notnull().any():
                st.markdown("##### 🗺️ 후보지 지도")
                map_df = df_sites[['lat', 'lon']].copy()
                st.map(map_df, size=40, zoom=11)
    else:
        st.warning("pandas가 설치되지 않아 후보지 목록을 표시할 수 없습니다.")
        st.info("`pip install pandas`를 실행한 뒤 다시 시도해주세요.")

def render_phase1_3(project_name, location, project_goals, additional_info):
    st.markdown("### Mission 1 · Phase 1.3 — 프로그램 확정")
    
    # Step 1: 요구사항 불러오기
    with st.expander("📄 1단계 · 요구사항 불러오기", expanded=not st.session_state.get('phase1_3_requirements_loaded')):
        st.caption("Phase 1.1에서 자동으로 불러오거나, 직접 요구사항 텍스트를 붙여넣으세요.")
        
        col_req_source = st.columns([1, 1])
        with col_req_source[0]:
            if st.session_state.get('phase1_requirements_structured'):
                st.success(" Phase 1.1 요구사항이 있습니다.")
                if st.button("Phase 1.1 요구사항 사용", key="phase1_3_use_phase1_requirements"):
                    st.session_state['phase1_3_requirements_text'] = st.session_state['phase1_requirements_structured']
                    st.session_state['phase1_3_requirements_loaded'] = True
                    st.rerun()
            else:
                st.info("Phase 1.1을 먼저 완료하거나 오른쪽에서 직접 입력하세요.")
        
        with col_req_source[1]:
            manual_req_text = st.text_area(
                "요구사항 텍스트 직접 입력",
                height=150,
                placeholder="요구사항 텍스트를 붙여넣으세요...",
                key="phase1_3_manual_requirements_input"
            )
            if st.button("입력한 요구사항 사용", key="phase1_3_use_manual_requirements"):
                if manual_req_text.strip():
                    st.session_state['phase1_3_requirements_text'] = manual_req_text
                    st.session_state['phase1_3_requirements_loaded'] = True
                    st.success("요구사항을 불러왔습니다.")
                    st.rerun()
                else:
                    st.warning("요구사항 텍스트를 입력해주세요.")
        
        if st.session_state.get('phase1_3_requirements_text'):
            st.markdown("##### 불러온 요구사항")
            st.text_area(
                "현재 요구사항",
                value=st.session_state['phase1_3_requirements_text'],
                height=200,
                disabled=True,
                key="phase1_3_requirements_display"
            )
    
    if not st.session_state.get('phase1_3_requirements_loaded'):
        st.info("⬆️ 요구사항을 먼저 불러와주세요.")
        return
    
    requirements_text = st.session_state.get('phase1_3_requirements_text', '')
    
    # Step 2: 후보지 선택
    with st.expander("📍 2단계 · 후보지 선택", expanded=not st.session_state.get('phase1_3_selected_site')):
        st.caption("Phase 1.2에서 검토한 후보지 중 하나를 선택하세요.")
        
        candidate_sites = st.session_state.get('phase1_candidate_sites', [])
        if not candidate_sites:
            st.warning("Phase 1.2에서 후보지를 먼저 입력해주세요.")
        else:
            site_options = [site.get('name', f"후보지 {i+1}") for i, site in enumerate(candidate_sites)]
            selected_site_name = st.selectbox(
                "최종 선택 후보지",
                options=site_options,
                key="phase1_3_site_selector"
            )
            
            if st.button("선택 확정", key="phase1_3_confirm_site"):
                selected_idx = site_options.index(selected_site_name)
                st.session_state['phase1_3_selected_site'] = candidate_sites[selected_idx]
                st.session_state['phase1_3_selected_site_name'] = selected_site_name
                st.success(f"'{selected_site_name}'을(를) 선택했습니다.")
                st.rerun()
        
        if st.session_state.get('phase1_3_selected_site'):
            selected = st.session_state['phase1_3_selected_site']
            st.info(f"**선택된 후보지**: {st.session_state.get('phase1_3_selected_site_name', 'N/A')}")
            col_site_info = st.columns(3)
            col_site_info[0].metric("면적", f"{selected.get('area_m2', 0):,.0f}㎡")
            col_site_info[1].metric("경사도", f"{selected.get('slope_percent', 0):.1f}%")
            col_site_info[2].metric("총점", f"{selected.get('score', 0):.1f}점")
    
    if not st.session_state.get('phase1_3_selected_site'):
        st.info("⬆️ 후보지를 먼저 선택해주세요.")
        return
    
    # Step 3: AI 블록 실행 (블록 5, 6, 7 제거됨)
    # 제거된 블록들: phase1_facility_program, phase1_facility_area_reference, phase1_facility_area_calculation
    # 관련 섹션 제거됨



# 메인 컨텐츠
tab_project = tab_blocks = tab_run = tab_download = None  # type: ignore
tab_project, tab_blocks, tab_run, tab_download = st.tabs(
    ["기본 정보 & 파일 업로드", "분석 블록 선택", "분석 실행", "결과 다운로드"]
)

project_name = st.session_state.get("project_name", "")
location = st.session_state.get("location", "")
project_goals = st.session_state.get("project_goals", "")
additional_info = st.session_state.get("additional_info", "")

with tab_project:
    st.header("프로젝트 기본 정보 입력")
    st.caption("프로젝트 기본 정보는 이 탭에서 별도로 관리됩니다. 입력값은 자동 저장됩니다.")

    # 디버그 정보 (개발 중 확인용) - 숨김 처리
    # with st.expander("🔍 세션 상태 확인 (디버그)", expanded=False):
    #     st.caption("현재 세션에 저장된 프로젝트 정보를 확인할 수 있습니다.")
    #     debug_info = {
    #         "프로젝트명": st.session_state.get('project_name', '(없음)'),
    #         "위치": st.session_state.get('location', '(없음)'),
    #         "위도": st.session_state.get('latitude', '(없음)'),
    #         "경도": st.session_state.get('longitude', '(없음)'),
    #         "프로젝트 목표": st.session_state.get('project_goals', '(없음)')[:50] + "..." if len(st.session_state.get('project_goals', '')) > 50 else st.session_state.get('project_goals', '(없음)'),
    #     }
    #     for key, value in debug_info.items():
    #         st.text(f"{key}: {value}")
    #
    #     # DB 저장 확인
    #     if 'pms_current_user' in st.session_state:
    #         user_id = st.session_state.pms_current_user.get('id')
    #         st.text(f"사용자 ID: {user_id}")
    #     else:
    #         st.warning("로그인 정보 없음")

    st.text_input(
        "프로젝트명",
        value=st.session_state.get("project_name", ""),
        placeholder="예: 삼척 스포츠아카데미",
        key="project_name"
    )
    # 지도에서 필지 선택 시 자동 주입된 경우 안내
    if st.session_state.get("_map_parcel_loaded") and st.session_state.get("location"):
        st.caption("📍 지도(필지 선택) 페이지에서 선택한 필지 주소가 자동으로 입력되었습니다.")
    st.text_input(
        "위치/지역",
        value=st.session_state.get("location", ""),
        placeholder="예: 강원도 삼척시 도계읍 일대 (또는 지도 페이지에서 필지 선택 시 자동 입력)",
        key="location"
    )
    
    st.text_area(
        "프로젝트 목표",
        value=st.session_state.get("project_goals", ""),
        placeholder="예: 국제 스포츠 아카데미 조성, 지역 경제 활성화, 교육·훈련 통합 프로그램 구축 등",
        height=80,
        key="project_goals"
    )
    st.text_area(
        "추가 정보",
        value=st.session_state.get("additional_info", ""),
        placeholder="특별한 제약조건이나 참고 사항이 있다면 입력하세요.",
        height=80,
        key="additional_info"
    )

    # 프로젝트 정보 저장/불러오기 버튼
    col_load, col_save = st.columns(2)

    with col_load:
        if st.button("📥 저장된 정보 불러오기", use_container_width=True, key="load_project_info"):
            # 불러오기 플래그 설정하고 즉시 rerun (위젯이 생성되기 전에)
            st.session_state['_force_load_session'] = True
            st.rerun()

    with col_save:
        save_button_clicked = st.button("✅ 프로젝트 정보 저장", use_container_width=True, type="primary", key="save_project_info")

    if save_button_clicked:
        # 세션 저장
        try:
            from auth.session_init import save_work_session, save_analysis_progress
            from database.db_manager import execute_query
            import json
            from datetime import datetime

            # 로그인 확인
            if 'pms_current_user' not in st.session_state:
                st.error("❌ 로그인 정보가 없습니다. 다시 로그인해주세요.")
                st.stop()

            user_id = st.session_state.pms_current_user.get('id')
            if not user_id:
                st.error("❌ 사용자 ID를 가져올 수 없습니다.")
                st.stop()

            # 현재 세션 상태 출력 (디버그)
            print(f"[저장] 사용자 ID: {user_id}")
            print(f"[저장] project_name: '{st.session_state.get('project_name')}'")
            print(f"[저장] location: '{st.session_state.get('location')}'")
            print(f"[저장] project_goals: '{st.session_state.get('project_goals', '')[:50]}...'")

            # 저장 실행
            save_work_session()
            save_analysis_progress(force=True)

            # 저장 확인 (디버그)
            check_result = execute_query(
                "SELECT session_data FROM analysis_sessions WHERE user_id = ? ORDER BY created_at DESC LIMIT 1",
                (user_id,)
            )
            if check_result:
                _raw = check_result[0]['session_data']
                saved_data = json.loads(_raw) if isinstance(_raw, str) else _raw
                print(f"[저장 확인] DB에 저장된 project_name: '{saved_data.get('project_name')}'")
                print(f"[저장 확인] DB에 저장된 location: '{saved_data.get('location')}'")

                # UI에 저장된 내용 표시
                st.success("✅ 프로젝트 정보가 저장되었습니다!")
                with st.expander("저장된 내용 확인", expanded=True):
                    st.write(f"**프로젝트명**: {saved_data.get('project_name', '(없음)')}")
                    st.write(f"**위치**: {saved_data.get('location', '(없음)')}")
                    st.write(f"**총 {len(saved_data)}개 항목 저장됨**")
            else:
                st.warning("⚠️ 저장은 완료되었으나 확인할 수 없습니다.")

        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"[저장 오류 전체 내역]:\n{error_details}")
            st.error(f"❌ 저장 실패: {str(e)}")
            with st.expander("오류 상세 정보"):
                st.code(error_details)

    st.markdown("---")
    st.header("파일 업로드")

    # ── Supabase에 저장된 파일 목록/다운로드 ────────────────────────────────
    try:
        from auth.file_storage import get_project_files, download_project_file
        _pid_files = st.session_state.get("current_project_id")
        if _pid_files:
            saved_files = get_project_files(_pid_files)
            if saved_files:
                with st.expander("☁️ 저장된 업로드 파일", expanded=False):
                    st.caption("이 프로젝트에 저장된 파일 목록입니다. 필요 시 다시 다운로드할 수 있습니다.")
                    for f in saved_files[:20]:
                        fname = f.get("filename", "file")
                        spath = f.get("storage_path")
                        fsize = f.get("file_size_bytes") or 0
                        cols = st.columns([3, 1, 1])
                        with cols[0]:
                            st.write(f"**{fname}**")
                            st.caption(f"type={f.get('file_type')}, size={(fsize/1024/1024):.2f}MB")
                        with cols[1]:
                            if st.button("⬇️ 가져오기", key=f"pull_file_{f.get('id')}"):
                                if spath:
                                    data = download_project_file(spath)
                                    if data:
                                        st.session_state["pdf_uploaded"] = True
                                        st.session_state["file_storage_path"] = spath
                                        st.info("파일 바이너리를 내려받았습니다. 아래에서 다시 분석하거나 다운로드하세요.")
                                else:
                                    st.warning("storage_path가 없습니다.")
                        with cols[2]:
                            if spath:
                                data = None
                                try:
                                    data = download_project_file(spath)
                                except Exception:
                                    data = None
                                st.download_button(
                                    "다운로드",
                                    data=data,
                                    file_name=fname,
                                    key=f"dl_file_{f.get('id')}",
                                )
    except Exception as _files_err:
        print(f"[FileStorage] 저장 파일 목록 UI 오류: {_files_err}")
    
    uploaded_file = st.file_uploader(
        "파일을 업로드하세요",
        type=['pdf', 'docx', 'xlsx', 'xls', 'csv', 'txt', 'json', 'png', 'jpg', 'jpeg', 'webp'],
        help="도시 프로젝트 관련 문서를 업로드하세요 (PDF, Word, Excel, CSV, 텍스트, JSON 지원)"
    )
    
    if uploaded_file is not None:
        st.success(f"파일 업로드 완료: {uploaded_file.name}")
        
        # 파일 확장자 확인
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_bytes = uploaded_file.getvalue()

        # 이미지 파일: Gemini Vision으로 내용 읽기 → 텍스트 컨텍스트로 저장
        if file_extension in ["png", "jpg", "jpeg", "webp"]:
            try:
                from google import genai
                from google.genai import types
                from pdf_analyzer import _get_gemini_api_key

                api_key = _get_gemini_api_key()
                if not api_key:
                    st.error("이미지 읽기를 위해 `GEMINI_API_KEY`가 필요합니다. (설정/환경변수 또는 UI 키)")
                else:
                    client = genai.Client(api_key=api_key)
                    prompt = (
                        "이 이미지를 '도시/건축 프로젝트 분석' 관점에서 읽고, "
                        "보이는 핵심 요소(텍스트/도면/표/지도/다이어그램)를 구조화해 한국어로 요약해줘.\n\n"
                        "출력 형식:\n"
                        "1) 한줄 요약\n"
                        "2) 관찰된 요소(불릿)\n"
                        "3) 이미지 내 텍스트(OCR 느낌으로 최대한)\n"
                        "4) 분석에 유용한 키워드(10개)\n"
                    )
                    with st.spinner("🖼️ 이미지 내용 읽는 중(Gemini Vision)..."):
                        resp = client.models.generate_content(
                            model="gemini-2.5-flash",
                            contents=[
                                types.Content(
                                    role="user",
                                    parts=[
                                        types.Part.from_text(prompt),
                                        types.Part.from_bytes(data=file_bytes, mime_type=uploaded_file.type or "image/png"),
                                    ],
                                )
                            ],
                        )
                    text = (getattr(resp, "text", None) or "").strip()
                    if text:
                        # 기존 변수명을 유지: pdf_text에 이미지 설명을 넣어 이후 블록 분석에 바로 사용 가능
                        st.session_state["pdf_text"] = text
                        st.session_state["pdf_uploaded"] = True
                        st.session_state["file_type"] = "image"
                        st.session_state["file_analysis"] = {
                            "success": True,
                            "file_type": "image",
                            "text": text,
                            "char_count": len(text),
                            "word_count": len(text.split()),
                            "preview": text[:500] + "..." if len(text) > 500 else text,
                        }
                        st.session_state["uploaded_file"] = uploaded_file
                        st.success("이미지 읽기 완료! 추출된 텍스트를 분석에 사용합니다.")
                        with st.expander("이미지 읽기 결과(미리보기)"):
                            st.text(st.session_state["file_analysis"]["preview"])

                        # Storage 업로드는 기존 로직 그대로 수행하도록 아래로 계속 진행
                    else:
                        st.error("이미지 읽기 결과가 비어 있습니다.")
            except Exception as _img_err:
                st.error(f"이미지 읽기 실패: {_img_err}")
            # 이미지인 경우도 Storage 업로드/메타 저장을 위해 아래 로직은 계속 진행
        
        # 메모리에서 직접 파일 분석 (임시 파일 생성 없음)
        file_analyzer = UniversalFileAnalyzer()
        
        # 파일 분석 (메모리 기반)
        with st.spinner(f"{file_extension.upper()} 파일 분석 중..."):
            analysis_result = file_analyzer.analyze_file_from_bytes(
                file_bytes, 
                file_extension, 
                uploaded_file.name
            )
            
        if analysis_result['success']:
            st.success(f"{file_extension.upper()} 파일 분석 완료!")
            
            # 파일 정보 표시 (파일 크기는 업로드된 파일에서 직접 계산)
            file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
            st.info(f"파일 정보: {file_size_mb:.2f}MB, {analysis_result['word_count']}단어, {analysis_result['char_count']}문자")
            if analysis_result.get('truncated'):
                orig = analysis_result.get('original_char_count', 0)
                st.warning(f"파일이 너무 커서 앞부분 {analysis_result['char_count']:,}자만 분석에 사용됩니다. (원본: {orig:,}자)")
            
            # 파일 형식별 특별 정보 표시
            if analysis_result['file_type'] == 'excel':
                st.info(f"Excel 시트: {', '.join(analysis_result['sheet_names'])} ({analysis_result['sheet_count']}개 시트)")
            elif analysis_result['file_type'] == 'csv':
                enc = analysis_result.get('encoding', 'utf-8')
                st.info(f"CSV 데이터: {analysis_result['shape'][0]}행 × {analysis_result['shape'][1]}열 | 인코딩: {enc}")
            elif analysis_result['file_type'] == 'pdf':
                method = analysis_result.get('method', 'pymupdf')
                quality = analysis_result.get('quality_score', '-')
                st.info(f"PDF 추출 방법: {method} | 품질 점수: {quality}/100")
                if analysis_result.get('is_scanned'):
                    st.warning("스캔된 PDF로 감지되었습니다. Gemini API 키가 있으면 자동으로 OCR 처리됩니다.")
            elif analysis_result['file_type'] == 'docx':
                h = analysis_result.get('heading_count', 0)
                t = analysis_result.get('table_count', 0)
                st.info(f"Word 문서: 헤딩 {h}개, 표 {t}개")
            elif analysis_result['file_type'] == 'json':
                if analysis_result.get('summarized'):
                    st.info("JSON 파일이 크기가 커서 구조 요약 모드로 변환되었습니다. (키 목록 + 샘플 항목)")
            
            # 세션에 저장
            st.session_state['pdf_text'] = analysis_result['text']  # 기존 변수명 유지
            st.session_state['pdf_uploaded'] = True
            st.session_state['file_type'] = analysis_result['file_type']
            st.session_state['file_analysis'] = analysis_result
            st.session_state['uploaded_file'] = uploaded_file  # 파일 객체 저장

            # ── Supabase Storage 업로드 ────────────────────────────────────────
            try:
                from auth.file_storage import upload_project_file, save_file_meta
                from auth.project_manager import get_or_create_current_project
                _uid_fs = st.session_state.pms_current_user.get('id') if st.session_state.get('pms_current_user') else None
                if _uid_fs:
                    _pid_fs = get_or_create_current_project(_uid_fs)
                    _storage_path = upload_project_file(
                        user_id=_uid_fs,
                        project_id=_pid_fs,
                        filename=uploaded_file.name,
                        file_bytes=file_bytes,
                    )
                    if _storage_path:
                        st.session_state['file_storage_path'] = _storage_path
                        save_file_meta(
                            project_id=_pid_fs,
                            user_id=_uid_fs,
                            filename=uploaded_file.name,
                            file_type=analysis_result['file_type'],
                            storage_path=_storage_path,
                            char_count=analysis_result.get('char_count', 0),
                            file_size_bytes=len(file_bytes),
                            file_meta={
                                'quality_score': analysis_result.get('quality_score'),
                                'method': analysis_result.get('method'),
                            },
                        )
                        st.caption(f"☁️ 파일이 저장소에 업로드되었습니다.")
            except Exception as _fs_err:
                print(f"[FileStorage] 업로드 오류: {_fs_err}")

            # 텍스트 미리보기
            with st.expander(f"{file_extension.upper()} 내용 미리보기"):
                st.text(analysis_result['preview'])

            # 파일 업로드 확인 버튼
            if st.button("✅ 파일 분석 완료 확인", use_container_width=True, type="primary", key="confirm_file_upload"):
                try:
                    from auth.session_init import save_work_session, save_analysis_progress
                    save_work_session()
                    save_analysis_progress(force=True)
                    st.success("파일 분석 결과가 저장되었습니다! '분석 블록 선택' 탭으로 이동하세요.")
                except Exception as e:
                    st.warning(f"저장 중 오류: {e}")
                    st.success("파일 분석이 확인되었습니다. '분석 블록 선택' 탭으로 이동하세요.")
        else:
            st.error(f"{file_extension.upper()} 파일 분석에 실패했습니다: {analysis_result.get('error', '알 수 없는 오류')}")

    # 입력값 최신화
    project_name = st.session_state.get("project_name", "")
    location = st.session_state.get("location", "")
    project_goals = st.session_state.get("project_goals", "")
    additional_info = st.session_state.get("additional_info", "")

with tab_blocks:
    st.header("분석 블록 선택")
    
    # 기본 정보나 파일 중 하나라도 있으면 진행
    has_basic_info = any([project_name, location, project_goals, additional_info])
    has_file = st.session_state.get('pdf_uploaded', False)
    
    if not has_basic_info and not has_file:
        st.warning("프로젝트 기본 정보를 입력하거나 파일을 업로드해주세요.")
        st.stop()


    # get_example_blocks()는 이미 모든 블록(custom 포함)을 반환하므로 중복 방지
    all_blocks = get_example_blocks()
    block_lookup = {
        block.get('id'): block
        for block in all_blocks
        if isinstance(block, dict) and block.get('id')
    }

    grouped_blocks = group_blocks_by_category(all_blocks)
    
    if not grouped_blocks:
        st.info("사용 가능한 분석 블록이 없습니다.")
    else:
        st.subheader("블록 목록")
        ordered_categories = iter_categories_in_order(grouped_blocks)
        total_categories = len(ordered_categories)
        
        for idx, category in enumerate(ordered_categories):
            st.markdown(f"#### 📂 {category}")
            for block_idx, block in enumerate(grouped_blocks[category]):
                block_id = block.get('id')
                if not block_id:
                    continue
                
                is_custom_block = (
                    block.get('created_by') == 'user' or
                    str(block_id).startswith('custom_')
                )
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    # 공간 데이터 연동 여부 확인
                    block_spatial = st.session_state.get('block_spatial_data', {})
                    is_linked = block_id in block_spatial

                    block_name = block.get('name', '이름 없음')

                    # 사용자 블록에 [개인]/[팀] 태그 추가
                    if is_custom_block or block.get('_db_id'):
                        visibility = block.get('_visibility', block.get('visibility', ''))
                        if visibility in ['personal', 'PERSONAL']:
                            if not block_name.startswith('[개인]'):
                                block_name = f"[개인] {block_name}"
                        elif visibility in ['team', 'TEAM']:
                            if not block_name.startswith('[팀]'):
                                block_name = f"[팀] {block_name}"

                    if is_linked:
                        linked_layer = block_spatial[block_id].get('layer_name', '')
                        st.markdown(f"**{block_name}** 📍")
                        st.caption(f"🔗 연동: {linked_layer}")
                    else:
                        st.markdown(f"**{block_name}**")

                    description = block.get('description')
                    if description:
                        st.caption(description)
                
                with col2:
                    is_selected = block_id in st.session_state['selected_blocks']
                    # 카테고리와 인덱스를 포함하여 고유한 key 생성
                    unique_key = f"select_{category}_{block_idx}_{block_id}"
                    checkbox_value = st.checkbox(
                        "선택",
                        key=unique_key,
                        value=is_selected
                    )
                    
                    if checkbox_value and not is_selected:
                        st.session_state['selected_blocks'].append(block_id)
                    elif not checkbox_value and is_selected:
                        # 분석 세션 진행 중이고 cot_plan에 있는 블록은 제거하지 않음
                        if st.session_state.get('cot_session') and block_id in st.session_state.get('cot_plan', []):
                            print(f"[DEBUG 체크박스] 블록 {block_id} 제거 방지 (cot_plan에 있음)")
                        else:
                            print(f"[DEBUG 체크박스] 블록 {block_id} 제거됨")
                            st.session_state['selected_blocks'].remove(block_id)
            
            if idx < total_categories - 1:
                st.divider()
    
    # 선택된 블록들 표시 및 순서 조정
    selected_blocks = st.session_state['selected_blocks']
    if selected_blocks:
        st.success(f"{len(selected_blocks)}개 블록이 선택되었습니다:")
        
        # 선택된 블록들의 정보를 DataFrame으로 구성
        import pandas as pd
        
        block_info_list = []
        for order, block_id in enumerate(selected_blocks, start=1):
            block = block_lookup.get(block_id)
            block_name = block.get('name', '알 수 없음') if block else "알 수 없음"

            # 사용자 블록에 [개인]/[팀] 태그 추가
            if block:
                is_custom = block.get('created_by') == 'user' or str(block_id).startswith('custom_') or block.get('_db_id')
                if is_custom:
                    visibility = block.get('_visibility', block.get('visibility', ''))
                    if visibility in ['personal', 'PERSONAL'] and not block_name.startswith('[개인]'):
                        block_name = f"[개인] {block_name}"
                    elif visibility in ['team', 'TEAM'] and not block_name.startswith('[팀]'):
                        block_name = f"[팀] {block_name}"

            block_description = block.get('description', '') if block else ""
            block_category = resolve_block_category(block) if block else "기타"
            block_info_list.append({
                '순서': order,
                '카테고리': block_category,
                '블록명': block_name,
                '설명': block_description,
                '블록ID': block_id
            })
        
        # 순서 조정을 위한 데이터프레임 생성
        df = pd.DataFrame(block_info_list)
        
        st.subheader("선택된 블록 목록 및 순서 조정")
        st.caption("💡 순서 컬럼의 숫자를 직접 수정하거나, 오른쪽에서 행을 선택하여 화살표 버튼으로 순서를 변경할 수 있습니다.")
        
        # 표와 버튼을 나란히 배치
        col_table, col_buttons = st.columns([5, 1])
        
        with col_table:
            # 수정 가능한 데이터 에디터로 순서 조정
            edited_df = st.data_editor(
                df[['순서', '카테고리', '블록명', '설명']],
                use_container_width=True,
                num_rows="fixed",
                key="block_order_editor",
                column_config={
                    "순서": st.column_config.NumberColumn(
                        "순서",
                        help="분석 실행 순서 (숫자가 작을수록 먼저 실행)",
                        min_value=1,
                        max_value=len(block_info_list),
                        step=1
                    ),
                    "카테고리": st.column_config.TextColumn(
                        "카테고리",
                        disabled=True
                    ),
                    "블록명": st.column_config.TextColumn(
                        "블록명",
                        disabled=True
                    ),
                    "설명": st.column_config.TextColumn(
                        "설명",
                        disabled=True
                    )
                }
            )

            # 순서 변경사항 감지 (세션에 저장)
            original_order = df['순서'].tolist()
            edited_order = edited_df['순서'].tolist()

            # 변경사항이 있는지 표시
            order_changed = original_order != edited_order
            if order_changed:
                # 중복 검사
                if len(set(edited_order)) != len(edited_order):
                    st.warning("⚠️ 순서 값이 중복되었습니다. 고유한 숫자를 입력해주세요.")
                else:
                    st.info("✏️ 순서가 변경되었습니다. 아래 '순서 적용' 버튼을 클릭하세요.")
                    # 변경된 순서를 임시로 저장
                    st.session_state['pending_block_order'] = edited_df

        with col_buttons:
            st.markdown("")  # 상단 여백
            st.markdown("")  # 상단 여백
            
            # 선택된 행 인덱스 초기화 및 유효성 검사
            if 'selected_block_row_index' not in st.session_state:
                st.session_state.selected_block_row_index = 0
            
            # 인덱스가 유효한 범위 내에 있는지 확인
            max_index = len(block_info_list) - 1
            if st.session_state.selected_block_row_index > max_index:
                st.session_state.selected_block_row_index = max_index
            if st.session_state.selected_block_row_index < 0:
                st.session_state.selected_block_row_index = 0
            
            # 행 선택을 위한 selectbox
            block_options = [f"{i+1}. {row['블록명']}" for i, row in df.iterrows()]
            selected_row_display = st.selectbox(
                "행 선택:",
                options=block_options,
                index=st.session_state.selected_block_row_index,
                key="block_row_selector",
                label_visibility="collapsed"
            )
            
            # 선택된 인덱스 업데이트
            selected_row_index = block_options.index(selected_row_display)
            st.session_state.selected_block_row_index = selected_row_index
            
            st.markdown("")  # 여백
            
            # 위/아래 화살표 버튼
            move_up_disabled = (selected_row_index == 0)
            if st.button("⬆️", key="move_block_up", disabled=move_up_disabled, use_container_width=True, help="위로 이동"):
                if selected_row_index > 0:
                    current_blocks = st.session_state['selected_blocks'].copy()
                    # 선택된 블록과 위 블록 교환
                    current_blocks[selected_row_index], current_blocks[selected_row_index - 1] = \
                        current_blocks[selected_row_index - 1], current_blocks[selected_row_index]
                    st.session_state['selected_blocks'] = current_blocks
                    st.session_state.selected_block_row_index = selected_row_index - 1
                    st.success("블록이 위로 이동되었습니다!")
                    st.rerun()
            
            move_down_disabled = (selected_row_index == len(st.session_state['selected_blocks']) - 1)
            if st.button("⬇️", key="move_block_down", disabled=move_down_disabled, use_container_width=True, help="아래로 이동"):
                if selected_row_index < len(st.session_state['selected_blocks']) - 1:
                    current_blocks = st.session_state['selected_blocks'].copy()
                    # 선택된 블록과 아래 블록 교환
                    current_blocks[selected_row_index], current_blocks[selected_row_index + 1] = \
                        current_blocks[selected_row_index + 1], current_blocks[selected_row_index]
                    st.session_state['selected_blocks'] = current_blocks
                    st.session_state.selected_block_row_index = selected_row_index + 1
                    st.success("블록이 아래로 이동되었습니다!")
                    st.rerun()

            # 순서 직접 수정 적용 버튼
            if 'pending_block_order' in st.session_state:
                st.markdown("")  # 여백
                if st.button("✅ 순서 적용", key="apply_block_order", type="primary", use_container_width=True, help="편집한 순서를 적용합니다"):
                    try:
                        pending_df = st.session_state['pending_block_order']
                        edited_order = pending_df['순서'].tolist()

                        # 중복 검사
                        if len(set(edited_order)) == len(edited_order):
                            sorted_indices = pending_df.sort_values('순서', kind="stable").index
                            new_blocks = [df.loc[idx, '블록ID'] for idx in sorted_indices]
                            st.session_state['selected_blocks'] = new_blocks
                            del st.session_state['pending_block_order']
                            st.success("블록 순서가 업데이트되었습니다!")
                            st.rerun()
                        else:
                            st.error("순서 값이 중복되었습니다.")
                    except Exception as e:
                        st.error(f"순서 업데이트 중 오류: {e}")

        # 공간 데이터 연동 섹션
        if st.session_state.get('downloaded_geo_data'):
            st.markdown("---")
            st.subheader("🔗 공간 데이터 연동")
            st.caption("Mapping에서 조회한 공간 데이터를 각 블록의 분석에 활용할 수 있습니다.")

            # block_spatial_data 초기화
            if 'block_spatial_data' not in st.session_state:
                st.session_state.block_spatial_data = {}

            # 사용 가능한 레이어 목록
            available_layers = list(st.session_state.downloaded_geo_data.keys())
            layer_info = {
                layer_name: f"{layer_name} ({data.get('feature_count', 0)}개)"
                for layer_name, data in st.session_state.downloaded_geo_data.items()
            }

            # 각 선택된 블록에 대해 레이어 연동 UI
            for block_id in selected_blocks:
                block = block_lookup.get(block_id)
                block_name = block.get('name', block_id) if block else block_id

                # 현재 연동된 레이어 가져오기
                current_linked = []
                if block_id in st.session_state.block_spatial_data:
                    current_linked = st.session_state.block_spatial_data[block_id].get('layers', [])

                col_block, col_layer = st.columns([2, 3])
                with col_block:
                    st.markdown(f"**{block_name}**")
                with col_layer:
                    # multiselect로 레이어 선택
                    selected_layers = st.multiselect(
                        "연동할 레이어",
                        options=available_layers,
                        default=[l for l in current_linked if l in available_layers],
                        format_func=lambda x: layer_info.get(x, x),
                        key=f"layer_link_{block_id}",
                        label_visibility="collapsed"
                    )

                    # 선택 변경 시 block_spatial_data 업데이트
                    if selected_layers:
                        combined_features = []
                        total_count = 0
                        for layer_name in selected_layers:
                            if layer_name in st.session_state.downloaded_geo_data:
                                data = st.session_state.downloaded_geo_data[layer_name]
                                geojson = data.get('geojson', {})
                                for feature in geojson.get('features', []):
                                    feature_copy = dict(feature)
                                    if 'properties' not in feature_copy:
                                        feature_copy['properties'] = {}
                                    feature_copy['properties']['_layer'] = layer_name
                                    combined_features.append(feature_copy)
                                total_count += data.get('feature_count', 0)

                        st.session_state.block_spatial_data[block_id] = {
                            'layer_name': ', '.join(selected_layers),
                            'geojson': {'type': 'FeatureCollection', 'features': combined_features},
                            'feature_count': total_count,
                            'layers': selected_layers
                        }
                    elif block_id in st.session_state.block_spatial_data:
                        # 선택 해제 시 삭제
                        del st.session_state.block_spatial_data[block_id]

            # 연동 현황 요약
            linked_blocks = [bid for bid in selected_blocks if bid in st.session_state.block_spatial_data]
            if linked_blocks:
                st.success(f"✓ {len(linked_blocks)}개 블록에 공간 데이터가 연동되었습니다.")

        # 블록 선택 완료 버튼
        st.markdown("---")
        if st.button("✅ 블록 선택 완료", use_container_width=True, type="primary", key="confirm_block_selection"):
            try:
                from auth.session_init import save_work_session, save_analysis_progress
                save_work_session()
                save_analysis_progress(force=True)
                st.success(f"{len(selected_blocks)}개 블록이 선택되었습니다! '분석 실행' 탭으로 이동하세요.")
            except Exception as e:
                st.warning(f"저장 중 오류: {e}")
                st.success(f"{len(selected_blocks)}개 블록 선택 완료! '분석 실행' 탭으로 이동하세요.")
    else:
        st.warning("분석할 블록을 선택해주세요.")

with tab_run:
    st.header("분석 실행")
    has_basic_info = any([project_name, location, project_goals, additional_info])
    has_file = st.session_state.get('pdf_uploaded', False)
    has_existing_results = bool(st.session_state.get('analysis_results') or st.session_state.get('cot_results'))

    # 분석 결과가 있으면 기본 정보 체크 스킵 (세션 복원 시)
    if not has_existing_results:
        if not has_basic_info and not has_file:
            st.warning("프로젝트 기본 정보를 입력하거나 파일을 업로드해주세요.")
            st.stop()

    selected_blocks = st.session_state.get('selected_blocks', [])
    if not selected_blocks and not has_existing_results:
        st.warning("먼저 분석 블록을 선택해주세요.")
        st.stop()

    # get_example_blocks()는 이미 모든 블록(custom 포함)을 반환하므로 중복 방지
    all_blocks = get_example_blocks()
    block_lookup = {
        block.get('id'): block
        for block in all_blocks
        if isinstance(block, dict) and block.get('id')
    }

    st.subheader("분석 대상 정보")
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**프로젝트 정보**")
        if project_name:
            st.write(f"• 프로젝트명: {project_name}")
        if location:
            st.write(f"• 위치/지역: {location}")
        if project_goals:
            ellipsis = "..." if len(project_goals) > 100 else ""
            st.write(f"• 프로젝트 목표: {project_goals[:100]}{ellipsis}")
        if additional_info:
            ellipsis = "..." if len(additional_info) > 100 else ""
            st.write(f"• 추가 정보: {additional_info[:100]}{ellipsis}")

    with col2:
        st.markdown("**파일 정보**")
        if has_file:
            file_analysis = st.session_state.get('file_analysis', {})
            file_name = "N/A"
            if st.session_state.get('uploaded_file'):
                file_name = st.session_state['uploaded_file'].name
            elif uploaded_file is not None:
                file_name = uploaded_file.name
            st.write(f"• 파일명: {file_name}")
            st.write(f"• 파일 유형: {file_analysis.get('file_type', 'N/A')}")
            st.write(f"• 텍스트 길이: {file_analysis.get('char_count', 0)}자")
            st.write(f"• 단어 수: {file_analysis.get('word_count', 0)}단어")
        else:
            st.write("• 파일 없음 (기본 정보만 사용)")
        reference_docs = st.session_state.get('reference_documents', [])
        if reference_docs:
            total_chars = sum(doc.get('char_count', 0) for doc in reference_docs)
            st.write(f"• 참고 자료: {len(reference_docs)}건 ({total_chars:,}자)")

    # Mapping 필지 정보 연동 상태 표시
    _sf = st.session_state.get('site_fields')
    _geo = st.session_state.get('downloaded_geo_data')
    if _sf or _geo:
        st.markdown("---")
        st.markdown("**🗺️ Mapping 연동 현황**")
        if _sf:
            _addr = _sf.get('site_address', '')
            _area = _sf.get('site_area', '')
            _zoning = _sf.get('zoning', '')
            st.success(f"✓ 필지 정보: {_addr}  {_area}  {_zoning}")
            with st.expander("상세 필지 정보", expanded=False):
                for _key, _label in [
                    ('land_category','지목'), ('parcel_count','필지 수'),
                    ('official_price_per_m2','공시지가'), ('land_restrictions','용도계획 제한'),
                    ('land_ownership','소유정보'), ('building_height','건물 높이'),
                    ('existing_building_purpose','기존 건물 용도'),
                    ('existing_building_area','기존 건물 연면적'),
                    ('existing_bcr','기존 건폐율'), ('existing_vlr','기존 용적률'),
                    ('nearby_buildings_count','주변 건물 수'),
                    ('nearby_building_uses','주변 건물 주요 용도'),
                ]:
                    _v = _sf.get(_key)
                    if _v and str(_v) not in ('True','False',''):
                        st.write(f"• {_label}: {_v}")
        if _geo:
            for _lname, _ldata in _geo.items():
                st.info(f"✓ 공간 레이어: {_lname} ({_ldata.get('feature_count', 0)}개 필지)")

    # 공간 데이터 연동 상태 표시
    if st.session_state.get('block_spatial_data'):
        st.markdown("---")
        st.markdown("**🔗 Mapping 블록 연동 데이터**")
        block_spatial_data = st.session_state.block_spatial_data
        linked_blocks = [bid for bid in selected_blocks if bid in block_spatial_data]
        if linked_blocks:
            for block_id in linked_blocks:
                spatial_info = block_spatial_data[block_id]
                st.success(f"✓ {block_id}: {spatial_info['layer_name']} ({spatial_info['feature_count']}개 피처)")
        else:
            st.info(" Mapping 페이지에서 블록에 공간 데이터를 연동할 수 있습니다.")

    st.markdown("---")

    base_text_candidates: List[str] = []
    if has_file:
        base_text_candidates.append(st.session_state.get('pdf_text', ''))
    reference_combined = st.session_state.get('reference_combined_text', '')
    if reference_combined:
        base_text_candidates.append(reference_combined)
    base_text_candidates.extend(filter(None, [project_name, location, project_goals, additional_info]))
    base_text_source = "\n\n".join([text for text in base_text_candidates if text]).strip()

    analysis_text = base_text_source

    project_info_payload = {
        "project_name": project_name,
        "location": location,
        "project_goals": project_goals,
        "additional_info": additional_info,
        "file_text": analysis_text
    }
    reference_docs_meta = st.session_state.get('reference_documents', [])
    reference_combined_text = st.session_state.get('reference_combined_text', '')
    if reference_docs_meta:
        project_info_payload["reference_documents"] = reference_docs_meta
    if reference_combined_text:
        project_info_payload["reference_text"] = reference_combined_text

    # 문서 요약 추가 (있는 경우)
    if st.session_state.get('document_summary'):
        project_info_payload["document_summary"] = st.session_state.document_summary

    # 위치 좌표 추가 (Google Maps용)
    if st.session_state.get('latitude') and st.session_state.get('longitude'):
        try:
            project_info_payload["latitude"] = float(st.session_state.latitude)
            project_info_payload["longitude"] = float(st.session_state.longitude)
        except (ValueError, TypeError):
            pass

    # Mapping 페이지에서 선택한 필지 정보 (site_fields) → site_context로 변환
    _site_fields = st.session_state.get('site_fields')
    if _site_fields:
        _field_labels = [
            ('site_address',              '주소'),
            ('site_area',                 '면적'),
            ('zoning',                    '용도지역/지구'),
            ('land_category',             '지목'),
            ('parcel_count',              '필지 수'),
            ('official_price_per_m2',     '공시지가'),
            ('land_restrictions',         '용도계획 제한'),
            ('land_ownership',            '소유정보'),
            ('building_height',           '건물 높이'),
            ('existing_building_purpose', '기존 건물 용도'),
            ('existing_building_area',    '기존 건물 연면적'),
            ('existing_bcr',              '기존 건폐율'),
            ('existing_vlr',              '기존 용적률'),
            ('existing_floors',           '기존 층수'),
            ('nearby_buildings_count',    '주변 건물 수'),
            ('nearby_building_uses',      '주변 건물 주요 용도'),
        ]
        _lines = ["### 📍 대상지 필지 현황 (VWorld API 확인 데이터)"]
        for _key, _label in _field_labels:
            _val = _site_fields.get(_key)
            if _val and str(_val) not in ('True', 'False', ''):
                _lines.append(f"- {_label}: {_val}")
        if _site_fields.get('nearby_buildings_summary'):
            _lines.append("")
            _lines.append(_site_fields['nearby_buildings_summary'])
        if len(_lines) > 1:
            project_info_payload['site_context'] = '\n'.join(_lines)

    spatial_notice = None
    try:
        spatial_contexts = []

        # 1. 업로드된 Shapefile 레이어
        if st.session_state.get('geo_layers') and len(st.session_state.geo_layers) > 0:
            from geo_data_loader import extract_spatial_context_for_ai
            for layer_name, layer_data in st.session_state.geo_layers.items():
                gdf = layer_data['gdf']
                layer_type = 'general'
                if any(keyword in layer_name for keyword in ['행정', '시군', '읍면', '법정', 'adm']):
                    layer_type = 'administrative'
                elif any(keyword in layer_name for keyword in ['공시', '가격', '지가', 'price']):
                    layer_type = 'land_price'
                elif any(keyword in layer_name for keyword in ['소유', '토지', 'owner']):
                    layer_type = 'ownership'
                spatial_text = extract_spatial_context_for_ai(gdf, layer_type)
                spatial_contexts.append(f"**레이어: {layer_name}**\n{spatial_text}")
        elif st.session_state.get('uploaded_gdf') is not None:
            from geo_data_loader import extract_spatial_context_for_ai
            gdf = st.session_state.uploaded_gdf
            layer_type = st.session_state.get('layer_type', 'general')
            spatial_text = extract_spatial_context_for_ai(gdf, layer_type)
            spatial_contexts.append(f"**업로드 레이어**\n{spatial_text}")

        # WFS 다운로드 데이터는 블록별로 선택되므로 여기서는 제외
        # (각 블록 실행 시점에 선택한 레이어가 feedback으로 추가됨)

        # 2. Mapping 페이지에서 블록에 연동된 공간 데이터
        if st.session_state.get('block_spatial_data'):
            block_spatial_data = st.session_state.block_spatial_data
            for block_id in selected_blocks:
                if block_id in block_spatial_data:
                    spatial_info = block_spatial_data[block_id]
                    layer_name = spatial_info['layer_name']
                    feature_count = spatial_info['feature_count']

                    # GeoJSON 요약 정보 추출
                    geojson = spatial_info.get('geojson', {})
                    features = geojson.get('features', [])

                    summary_text = f"**Mapping 연동 레이어: {layer_name}** (블록: {block_id})\n"
                    summary_text += f"- 총 피처 수: {feature_count}개\n"

                    # 속성별 분포 통계 계산 (용도지역, 건물용도 등)
                    if features:
                        from collections import Counter
                        # 용도지역/건물용도 관련 컬럼 찾기
                        zone_counters = {}
                        price_values = []
                        area_values = []

                        for feature in features:
                            props = feature.get('properties', {})
                            for key, value in props.items():
                                if value is None or value == '':
                                    continue
                                key_upper = key.upper()
                                # 용도지역 관련 컬럼
                                if any(k in key_upper for k in ['USG_NM', 'PRPOS_AREA_NM', '용도지역', 'ZONE_NM', 'JIJIMOK']):
                                    if '용도지역' not in zone_counters:
                                        zone_counters['용도지역'] = Counter()
                                    zone_counters['용도지역'][str(value)] += 1
                                # 건물용도 관련 컬럼
                                elif any(k in key_upper for k in ['PURPS_NM', 'MAIN_PURPS', '주용도', 'BDTYP_NM']):
                                    if '건물용도' not in zone_counters:
                                        zone_counters['건물용도'] = Counter()
                                    zone_counters['건물용도'][str(value)] += 1
                                # 공시지가
                                elif any(k in key_upper for k in ['PBLNTF', '공시지가', 'PRICE']):
                                    try:
                                        price_values.append(float(value))
                                    except:
                                        pass
                                # 면적
                                elif any(k in key_upper for k in ['AREA', '면적', 'LNDPCLR']):
                                    try:
                                        area_values.append(float(value))
                                    except:
                                        pass

                        # 분포 통계 텍스트 생성
                        for category, counter in zone_counters.items():
                            if counter:
                                summary_text += f"\n**{category} 분포:**\n"
                                for zone_name, count in counter.most_common(10):
                                    summary_text += f"  - {zone_name}: {count}개\n"

                        # 공시지가 통계
                        if price_values:
                            avg_price = sum(price_values) / len(price_values)
                            summary_text += f"\n**공시지가 통계:**\n"
                            summary_text += f"  - 평균: {int(avg_price):,}원/㎡\n"
                            summary_text += f"  - 최소: {int(min(price_values)):,}원/㎡\n"
                            summary_text += f"  - 최대: {int(max(price_values)):,}원/㎡\n"

                        # 면적 통계
                        if area_values:
                            total_area = sum(area_values)
                            avg_area = total_area / len(area_values)
                            summary_text += f"\n**면적 통계:**\n"
                            summary_text += f"  - 총 면적: {total_area:,.1f}㎡\n"
                            summary_text += f"  - 평균 면적: {avg_area:,.1f}㎡\n"

                    spatial_contexts.append(summary_text)

        # 공간 컨텍스트 통합 (업로드된 Shapefile만)
        if spatial_contexts:
            project_info_payload["spatial_data_context"] = "\n\n---\n\n".join(spatial_contexts)
            project_info_payload["has_geo_data"] = True
            spatial_notice = f"📍 {len(spatial_contexts)}개 공간 레이어 정보가 분석에 포함됩니다."
        else:
            project_info_payload["has_geo_data"] = False
    except Exception as e:
        st.warning(f"공간 데이터 통합 중 오류: {e}")
        project_info_payload["has_geo_data"] = False
    if spatial_notice:
        st.caption(spatial_notice)

    # 분석 세션이 비활성화 상태에서만 블록 불일치 시 초기화
    # (분석 중 블록 추가 시에는 초기화하지 않음)
    if st.session_state.cot_plan and st.session_state.cot_plan != selected_blocks and not st.session_state.cot_session:
        reset_step_analysis_state()

    st.markdown("### 단계별 분석 제어")
    control_col1, control_col2 = st.columns(2)
    with control_col1:
        if st.button("🔄 분석 세션 초기화", use_container_width=True):
            print("[DEBUG] 초기화 버튼 클릭됨")
            print(f"[DEBUG] 초기화 전 cot_results: {list(st.session_state.cot_results.keys())}")
            print(f"[DEBUG] 초기화 전 cot_current_index: {st.session_state.cot_current_index}")
            reset_step_analysis_state()
            print(f"[DEBUG] 초기화 후 cot_results: {list(st.session_state.cot_results.keys())}")
            print(f"[DEBUG] 초기화 후 cot_current_index: {st.session_state.cot_current_index}")
            st.success("분석 세션을 초기화했습니다.")
            st.rerun()
    prepare_disabled = not analysis_text
    with control_col2:
        if st.button("🚀 단계별 분석 세션 준비", type="primary", use_container_width=True, disabled=prepare_disabled):
            if not analysis_text:
                st.warning("분석에 사용할 텍스트가 없습니다.")
            else:
                try:
                    # analysis_runs/analysis_steps: 새 run 및 step row 생성
                    try:
                        from auth.project_manager import get_or_create_current_project
                        from database.analysis_steps_manager import create_run, create_steps
                        _u = st.session_state.get("pms_current_user") or {}
                        _uid = _u.get("id")
                        if _uid:
                            _pid = get_or_create_current_project(_uid)
                            input_snapshot = {
                                "project_name": st.session_state.get("project_name", ""),
                                "location": st.session_state.get("location", ""),
                                "project_goals": st.session_state.get("project_goals", ""),
                                "additional_info": st.session_state.get("additional_info", ""),
                                "file_type": st.session_state.get("file_type", ""),
                                "file_storage_path": st.session_state.get("file_storage_path", ""),
                                "selected_blocks": selected_blocks,
                            }
                            run_id = create_run(_uid, _pid, input_snapshot=input_snapshot)
                            if run_id:
                                st.session_state["current_analysis_run_id"] = run_id
                                # blocks payload는 id/name만 필요
                                _blocks_payload = []
                                for bid in selected_blocks:
                                    b = block_lookup.get(bid, {"id": bid, "name": bid})
                                    _blocks_payload.append({"id": bid, "name": b.get("name", bid)})
                                _step_map = create_steps(run_id, _pid, _uid, _blocks_payload)
                                st.session_state["analysis_step_id_map"] = _step_map
                    except Exception as _init_steps_err:
                        print(f"[AnalysisSteps] run/steps 생성 실패: {_init_steps_err}")

                    # 세션 준비 시 모든 이전 상태를 완전히 초기화
                    EnhancedArchAnalyzer.reset_lm()
                    st.session_state.pop('cot_analyzer', None)
                    st.session_state.pop('_last_analyzer_provider', None)
                    
                    # 이전 세션 완전히 제거
                    st.session_state.cot_session = None
                    st.session_state.cot_plan = []
                    st.session_state.cot_current_index = 0
                    st.session_state.cot_results = {}
                    st.session_state.cot_progress_messages = []
                    st.session_state.cot_history = []
                    st.session_state.cot_citations = {}
                    st.session_state.cot_feedback_inputs = {}
                    st.session_state.cot_running_block = None
                    st.session_state.skipped_blocks = []  # 건너뛴 블록 목록 초기화

                    analyzer = get_cot_analyzer()
                    if analyzer is None:
                        st.error("분석기를 초기화할 수 없습니다. 위의 오류 메시지를 확인하세요.")
                        st.stop()

                    # 문서 요약 생성 (충분한 텍스트가 있고, 아직 생성되지 않은 경우)
                    if analysis_text and len(analysis_text) > 500 and not st.session_state.get('document_summary'):
                        with st.spinner("📄 문서 요약 생성 중..."):
                            summary_result = analyzer.generate_document_summary(analysis_text)
                            if summary_result.get('success'):
                                st.session_state.document_summary = summary_result
                                doc_type = summary_result.get('document_type', '미확인')
                                key_topics_count = len(summary_result.get('key_topics', []))
                                st.info(f"✅ 문서 요약 완료: {doc_type} (핵심 키워드 {key_topics_count}개 추출)")
                            else:
                                st.warning(f"문서 요약 생성 실패: {summary_result.get('error', '알 수 없는 오류')}")

                    # RAG 시스템 구축 (블록별 컨텍스트 분리를 위해)
                    if analysis_text and len(analysis_text) > 200 and not st.session_state.get('doc_rag_system'):
                        try:
                            from rag_helper import build_rag_system_for_documents
                            with st.spinner("🔍 문서 인덱싱 중 (블록별 최적 컨텍스트 준비)..."):
                                rag_system = build_rag_system_for_documents(
                                    documents=[analysis_text],
                                    chunk_size=800,
                                    overlap=150
                                )
                                if rag_system.get("num_chunks", 0) > 0:
                                    st.session_state.doc_rag_system = rag_system
                                    st.info(f"✅ 문서 인덱싱 완료: {rag_system['num_chunks']}개 청크")
                        except Exception as _rag_err:
                            print(f"[RAG] 인덱싱 실패 (전체 문서로 폴백): {_rag_err}")

                    # Phase 4: 도시 지표 추출 및 정합성 검증
                    if analysis_text and not st.session_state.get('urban_indicator_results'):
                        try:
                            from utils.urban_indicators import UrbanIndicatorExtractor
                            _extractor = UrbanIndicatorExtractor()
                            _indicators = _extractor.extract(analysis_text)
                            if _indicators:
                                _validation = _extractor.validate(_indicators)
                                st.session_state.urban_indicator_results = {
                                    'indicators': _indicators,
                                    'validation': _validation,
                                }
                        except Exception as _ind_err:
                            print(f"[UrbanIndicators] 추출 실패: {_ind_err}")

                    # document_summary를 project_info_payload에 추가
                    if st.session_state.get('document_summary'):
                        project_info_payload['document_summary'] = st.session_state.document_summary

                    # 완전히 새로운 세션 생성 (previous_results는 빈 딕셔너리로 시작)
                    session = analyzer.initialize_cot_session(project_info_payload, analysis_text, len(selected_blocks))
                    # 세션의 previous_results가 빈 딕셔너리인지 확인
                    if 'previous_results' in session:
                        session['previous_results'] = {}
                    if 'cot_history' in session:
                        session['cot_history'] = []
                    
                    st.session_state.cot_session = session
                    st.session_state.cot_plan = selected_blocks.copy()
                    st.session_state.cot_current_index = 0
                    st.session_state.cot_results = {}
                    st.session_state.cot_progress_messages = []
                    st.session_state.cot_history = []
                    st.session_state.analysis_results = {}
                    st.session_state.cot_citations = {}
                    st.session_state.cot_feedback_inputs = {}
                    st.success("단계별 분석 세션이 준비되었습니다. 순서대로 블록을 실행하세요.")
                    st.rerun()
                except Exception as e:
                    st.error(f"분석기 초기화 실패: {e}")

    # Phase 4: 도시 지표 검증 결과 표시
    _ind_data = st.session_state.get('urban_indicator_results')
    if _ind_data and _ind_data.get('validation'):
        with st.expander("🏙️ 도시 지표 검증 결과", expanded=False):
            for v in _ind_data['validation']:
                icon = "✅" if v['ok'] else "⚠️"
                stated_str = f" (기재: {v['stated']})" if v.get('stated') is not None else ""
                st.markdown(
                    f"{icon} **{v['item']}**: {v['calculated']} {v['unit']}{stated_str} — {v['note']}"
                )

    active_plan = st.session_state.cot_plan if st.session_state.cot_session else selected_blocks

    # 분석 중 블록 추가 기능 (분석 실행 중일 때는 완전히 비활성화)
    is_analysis_running = st.session_state.get('cot_running_block') is not None

    # 분석 실행 중에는 블록 추가 UI를 전혀 렌더링하지 않음
    if not is_analysis_running and st.session_state.cot_session and st.session_state.cot_plan:
        with st.expander("➕ 블록 추가 (분석 진행 중)", expanded=False):
            st.caption("분석 세션이 진행 중일 때 새 블록을 추가할 수 있습니다.")

            # 현재 플랜에 없는 블록들만 표시
            current_plan_ids = set(st.session_state.cot_plan)
            available_to_add = [
                block for block in all_blocks
                if block.get('id') and block.get('id') not in current_plan_ids
            ]

            if available_to_add:
                # 블록 선택
                block_options = {block['id']: block.get('name', block['id']) for block in available_to_add}
                selected_block_to_add = st.selectbox(
                    "추가할 블록 선택",
                    options=list(block_options.keys()),
                    format_func=lambda x: block_options.get(x, x),
                    key="add_block_selector"
                )

                # 삽입 위치 선택
                insert_positions = ["현재 위치 (다음에 실행)", "플랜 마지막에 추가"]
                for i, plan_block_id in enumerate(st.session_state.cot_plan):
                    plan_block = block_lookup.get(plan_block_id, {})
                    plan_block_name = plan_block.get('name', plan_block_id)
                    insert_positions.append(f"{i+1}. {plan_block_name} 뒤에 삽입")

                insert_position = st.selectbox(
                    "삽입 위치",
                    options=insert_positions,
                    key="insert_position_selector"
                )

                if st.button("➕ 블록 추가", type="primary", key="add_block_btn"):
                    if selected_block_to_add:
                        print(f"[DEBUG 블록추가] 추가 전 cot_plan: {st.session_state.cot_plan}")
                        print(f"[DEBUG 블록추가] 추가할 블록: {selected_block_to_add}")
                        new_plan = st.session_state.cot_plan.copy()

                        adjust_index = False  # 인덱스 조정 필요 여부

                        if insert_position == "현재 위치 (다음에 실행)":
                            # 현재 인덱스에 삽입하고, 인덱스는 그대로 (새 블록이 바로 다음에 실행됨)
                            insert_idx = st.session_state.cot_current_index
                            adjust_index = False
                        elif insert_position == "플랜 마지막에 추가":
                            insert_idx = len(new_plan)
                            adjust_index = False
                        else:
                            # "N. 블록명 뒤에 삽입" 형식에서 인덱스 추출
                            try:
                                position_num = int(insert_position.split(".")[0])
                                insert_idx = position_num  # 해당 블록 뒤에 삽입
                                # 현재 인덱스보다 앞에 삽입되면 인덱스 조정 필요
                                adjust_index = (insert_idx <= st.session_state.cot_current_index)
                            except:
                                insert_idx = len(new_plan)
                                adjust_index = False

                        new_plan.insert(insert_idx, selected_block_to_add)
                        st.session_state.cot_plan = new_plan

                        # 인덱스 조정
                        if adjust_index:
                            st.session_state.cot_current_index += 1

                        # selected_blocks도 업데이트 (일관성 유지)
                        st.session_state.selected_blocks = new_plan.copy()

                        # 세션 저장 후 재시작
                        try:
                            from auth.session_init import save_work_session
                            save_work_session()
                        except Exception as e:
                            print(f"세션 저장 오류: {e}")

                        print(f"[DEBUG 블록추가] 추가 후 cot_plan: {st.session_state.cot_plan}")
                        print(f"[DEBUG 블록추가] 추가 후 selected_blocks: {st.session_state.selected_blocks}")
                        added_block = block_lookup.get(selected_block_to_add, {})
                        added_block_name = added_block.get('name', selected_block_to_add)
                        st.success(f"'{added_block_name}' 블록이 추가되었습니다.")
                        st.rerun()
            else:
                st.info("추가 가능한 블록이 없습니다. 모든 블록이 이미 플랜에 포함되어 있습니다.")

    st.markdown("### 단계 진행 현황")

    # DEBUG: 상태 확인 (콘솔에만 출력)
    print(f"[DEBUG] cot_session 존재: {st.session_state.cot_session is not None}")
    print(f"[DEBUG] cot_current_index: {st.session_state.cot_current_index}")
    print(f"[DEBUG] cot_results keys: {list(st.session_state.cot_results.keys())}")
    print(f"[DEBUG] cot_plan: {st.session_state.cot_plan}")
    if st.session_state.cot_session:
        print(f"[DEBUG] cot_session previous_results keys: {list(st.session_state.cot_session.get('previous_results', {}).keys())}")

    if not active_plan:
        st.info("분석 세션을 준비하면 단계별 진행 정보를 확인할 수 있습니다.")
    else:
        running_block = st.session_state.get('cot_running_block')
        skipped_blocks = st.session_state.get('skipped_blocks', [])
        for idx, block_id in enumerate(active_plan, start=1):
            block = block_lookup.get(block_id)
            block_name = block.get('name', block_id) if block else block_id
            category = resolve_block_category(block) if block else "기타"

            # 결과 확인 (cot_results와 analysis_results 둘 다 확인)
            has_result = (block_id in st.session_state.cot_results or
                         block_id in st.session_state.analysis_results)

            # 상태 배지 결정 (우선순위: 완료 > 진행중 > 건너뜀 > 대기 > 준비)
            if has_result:
                # 결과가 있으면 무조건 완료 (running_block보다 우선)
                status_badge = "✅ 완료"
            elif running_block == block_id:
                # 현재 실행 중인 블록
                status_badge = "⏳ 진행중"
            elif block_id in skipped_blocks:
                # 건너뛴 블록
                status_badge = "⏭️ 건너뜀"
            elif st.session_state.cot_session and idx == st.session_state.cot_current_index + 1:
                # 다음 실행 대상
                status_badge = "🟡 대기"
            else:
                # 준비 상태
                status_badge = "⚪ 준비"
            is_collapsed = status_badge in ["✅ 완료", "⏭️ 건너뜀"]
            expander = st.expander(f"{idx}. {block_name} · {status_badge}", expanded=(not is_collapsed))
            with expander:
                st.caption((block.get('description') if block else "설명이 없습니다.") or "설명이 없습니다.")

                # 네비게이션 버튼: 완료된 블록이나 대기/준비 상태 블록에서 이동 가능
                show_nav_button = False
                nav_button_label = ""

                if has_result:
                    # 완료된 블록: 재시작 버튼
                    show_nav_button = True
                    nav_button_label = "🔄 이 블록부터 재시작"
                elif st.session_state.cot_session and idx - 1 < st.session_state.cot_current_index:
                    # 현재 위치보다 이전 블록: 돌아가기 버튼
                    show_nav_button = True
                    nav_button_label = "⬅️ 이 블록으로 돌아가기"

                if show_nav_button:
                    col_nav, col_empty = st.columns([1, 2])
                    with col_nav:
                        if st.button(nav_button_label, key=f"nav_to_{block_id}", use_container_width=True):
                            # 현재 인덱스를 이 블록의 인덱스로 설정
                            st.session_state.cot_current_index = idx - 1  # 0-based index
                            # 이 블록과 이후 블록의 결과 삭제
                            blocks_to_remove = active_plan[idx - 1:]
                            for bid in blocks_to_remove:
                                if bid in st.session_state.cot_results:
                                    del st.session_state.cot_results[bid]
                                if bid in st.session_state.get('cot_citations', {}):
                                    del st.session_state.cot_citations[bid]
                            st.success(f"'{block_name}'(으)로 이동했습니다.")
                            st.rerun()

                if has_result:
                    # 피드백 유형 선택
                    from dspy_analyzer import FEEDBACK_TYPES
                    feedback_type_options = {
                        'auto': '자동 감지',
                        **{k: v['name'] for k, v in FEEDBACK_TYPES.items()}
                    }
                    feedback_type_key = f"feedback_type_{block_id}"
                    if feedback_type_key not in st.session_state:
                        st.session_state[feedback_type_key] = 'auto'

                    col_type, col_hint = st.columns([1, 2])
                    with col_type:
                        selected_feedback_type = st.selectbox(
                            "피드백 유형",
                            options=list(feedback_type_options.keys()),
                            format_func=lambda x: feedback_type_options[x],
                            key=feedback_type_key,
                            help="피드백 유형을 선택하면 AI가 해당 관점에서 재분석합니다."
                        )
                    with col_hint:
                        # 선택된 유형에 대한 힌트 표시
                        if selected_feedback_type != 'auto' and selected_feedback_type in FEEDBACK_TYPES:
                            hint_info = FEEDBACK_TYPES[selected_feedback_type]
                            st.caption(f"**{hint_info['description']}**")
                            st.caption(f"_{hint_info['hint']}_")

                    feedback_state_key = f"feedback_input_{block_id}"
                    if feedback_state_key not in st.session_state:
                        st.session_state[feedback_state_key] = st.session_state.cot_feedback_inputs.get(block_id, "")

                    # 유형별 placeholder 설정
                    placeholder_text = "재분석 시 반영할 메모, 수정 요청, 추가 지시사항을 입력하세요."
                    if selected_feedback_type != 'auto' and selected_feedback_type in FEEDBACK_TYPES:
                        placeholder_text = FEEDBACK_TYPES[selected_feedback_type]['hint']

                    feedback_text = st.text_area(
                        "피드백 입력",
                        key=feedback_state_key,
                        height=120,
                        placeholder=placeholder_text
                    )
                    st.session_state.cot_feedback_inputs[block_id] = feedback_text
                    rerun_disabled = st.session_state.cot_running_block is not None or not feedback_text.strip()
                    if st.button(
                        "피드백 반영 재분석",
                        key=f"rerun_btn_{block_id}",
                        disabled=rerun_disabled,
                        help="입력한 피드백을 반영하여 해당 블록만 다시 분석합니다."
                    ):
                        analyzer = get_cot_analyzer()
                        st.session_state.cot_running_block = block_id
                        rerun_step_index = active_plan.index(block_id) + 1 if block_id in active_plan else None
                        progress_placeholder = st.empty()
                        rerun_block_info = block or {"id": block_id, "name": block_id}

                        def rerun_progress(message: str) -> None:
                            progress_placeholder.info(message)

                        # 피드백 유형 전달 (auto이면 None)
                        actual_feedback_type = None if selected_feedback_type == 'auto' else selected_feedback_type

                        # analysis_steps 상태 업데이트(있으면)
                        try:
                            from database.analysis_steps_manager import set_step_status, save_step_payloads
                            step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                            sid = step_map.get(block_id)
                            if sid:
                                set_step_status(sid, "running")
                                _inputs = {
                                    "feedback": feedback_text.strip(),
                                    "feedback_type": actual_feedback_type,
                                }
                                save_step_payloads(sid, inputs=_inputs, outputs=None)
                        except Exception as _rerun_db_err:
                            print(f"[AnalysisSteps] rerun running 업데이트 실패: {_rerun_db_err}")

                        try:
                            with st.spinner("피드백 기반 재분석 중..."):
                                step_result = analyzer.run_cot_step(
                                    block_id,
                                    rerun_block_info,
                                    st.session_state.cot_session
                                    if st.session_state.cot_session
                                    else analyzer.initialize_cot_session(project_info_payload, analysis_text, len(active_plan)),
                                    progress_callback=rerun_progress,
                                    step_index=rerun_step_index,
                                    feedback=feedback_text.strip(),
                                    feedback_type=actual_feedback_type
                                )
                        finally:
                            st.session_state.cot_running_block = None

                        if step_result.get('success'):
                            st.session_state.cot_session = step_result['cot_session']
                            st.session_state.cot_results[block_id] = step_result['analysis']
                            analysis_result = step_result['analysis']
                            st.session_state.analysis_results[block_id] = analysis_result
                            # Citations 저장
                            if step_result.get('all_citations'):
                                st.session_state.cot_citations[block_id] = step_result['all_citations']

                            # analysis_steps 출력 저장(있으면)
                            try:
                                from database.analysis_steps_manager import set_step_status, save_step_payloads
                                step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                                sid = step_map.get(block_id)
                                if sid:
                                    outp = {
                                        "analysis": analysis_result,
                                        "citations": st.session_state.get("cot_citations", {}).get(block_id),
                                        "verifications": st.session_state.get("cot_verifications", {}).get(block_id),
                                    }
                                    save_step_payloads(sid, inputs=None, outputs=outp)
                                    set_step_status(sid, "completed")
                            except Exception as _rerun_save_err:
                                print(f"[AnalysisSteps] rerun step 저장 실패: {_rerun_save_err}")
                            
                            # 자동 저장
                            project_info = {
                                "project_name": st.session_state.get('project_name', ''),
                                "location": st.session_state.get('location', '')
                            }
                            save_analysis_result(block_id, analysis_result, project_info)

                            # 분석 진행 상태 실시간 저장
                            try:
                                from auth.session_init import save_analysis_progress
                                save_analysis_progress(force=True)
                            except Exception as e:
                                print(f"분석 진행 저장 오류: {e}")

                            st.session_state.cot_history = step_result['cot_session'].get('cot_history', st.session_state.cot_history)
                            st.success(f"{block_name} 블록을 피드백에 맞춰 재분석했습니다.")
                            st.rerun()
                        else:
                            try:
                                from database.analysis_steps_manager import set_step_status
                                step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                                sid = step_map.get(block_id)
                                if sid:
                                    set_step_status(sid, "failed", error=str(step_result.get("error", ""))[:800])
                            except Exception as _rerun_fail_err:
                                print(f"[AnalysisSteps] rerun failed 업데이트 실패: {_rerun_fail_err}")
                            st.error(f"재분석 실패: {step_result.get('error', '알 수 없는 오류')}")
                elif status_badge == "🟡 대기":
                    st.info("다음 실행 대상 블록입니다. 아래 버튼을 눌러 분석을 진행하세요.")

    if st.session_state.cot_session and st.session_state.cot_current_index < len(st.session_state.cot_plan):
        # 인덱스 유효성 검증 및 자동 조정
        # 현재 인덱스 앞의 블록들 중 완료되지 않은 블록이 있는지 확인
        completed_blocks = set(st.session_state.cot_results.keys()) | set(st.session_state.analysis_results.keys())
        uncompleted_before_current = []
        for i in range(st.session_state.cot_current_index):
            bid = st.session_state.cot_plan[i]
            if bid not in completed_blocks and bid not in st.session_state.get('skipped_blocks', []):
                uncompleted_before_current.append((i, bid))

        # 완료되지 않은 이전 블록이 있으면 인덱스를 첫 번째 미완료 블록으로 조정
        if uncompleted_before_current:
            first_uncompleted_idx, first_uncompleted_id = uncompleted_before_current[0]
            st.warning(f"⚠️ 이전 블록이 완료되지 않았습니다. {first_uncompleted_idx + 1}번째 블록으로 이동합니다.")
            st.session_state.cot_current_index = first_uncompleted_idx

        next_block_id = st.session_state.cot_plan[st.session_state.cot_current_index]
        next_block = block_lookup.get(next_block_id, {"id": next_block_id})
        next_block_name = next_block.get('name', next_block_id)

        # 다음 실행 대상 블록 명확히 표시
        st.info(f"🎯 다음 실행 대상: **{st.session_state.cot_current_index + 1}번째 블록 - {next_block_name}** (ID: `{next_block_id}`)")

        # 블록별 공간 데이터 선택 UI
        downloaded_geo_data = st.session_state.get('downloaded_geo_data', {})
        if downloaded_geo_data:
            with st.expander("🗺️ 이 블록에 공간 데이터 연결", expanded=False):
                st.caption("분석에 포함할 WFS 레이어를 선택하세요.")

                # 블록별 선택 상태 초기화
                if 'block_spatial_selection' not in st.session_state:
                    st.session_state.block_spatial_selection = {}

                # 현재 블록의 선택 상태
                current_selection = st.session_state.block_spatial_selection.get(next_block_id, [])

                # 레이어 선택 체크박스
                new_selection = []
                for layer_name, data in downloaded_geo_data.items():
                    is_checked = st.checkbox(
                        f"{layer_name} ({data['feature_count']}개)",
                        value=layer_name in current_selection,
                        key=f"spatial_block_{next_block_id}_{layer_name}"
                    )
                    if is_checked:
                        new_selection.append(layer_name)

                st.session_state.block_spatial_selection[next_block_id] = new_selection

                if new_selection:
                    st.success(f"선택: {len(new_selection)}개 레이어")
                else:
                    st.info("공간 데이터 없이 분석합니다.")

        # 실행, 멈춤, 건너뛰기 버튼
        is_running = st.session_state.cot_running_block is not None

        run_col, stop_col, skip_col = st.columns([3, 1, 1])
        with run_col:
            run_clicked = st.button(
                f"▶️ {st.session_state.cot_current_index + 1}단계 실행: {next_block_name}",
                type="primary",
                disabled=is_running,
                use_container_width=True
            )
        with stop_col:
            stop_clicked = st.button(
                "⏹️ 멈춤",
                disabled=not is_running,
                use_container_width=True,
                help="현재 실행 중인 분석을 중단합니다.",
                type="secondary"
            )
        with skip_col:
            skip_clicked = st.button(
                "⏭️ 건너뛰기",
                disabled=is_running,
                use_container_width=True,
                help="이 블록을 건너뛰고 다음 블록으로 진행합니다."
            )
        
        # 멈춤 처리
        if stop_clicked:
            st.session_state.cot_running_block = None
            st.warning(f"{next_block_name} 블록 분석을 중단했습니다. 페이지를 새로고침합니다.")
            # analysis_runs 취소 처리
            _run_id = st.session_state.get("current_analysis_run_id")
            if _run_id and not st.session_state.get(f"_run_finalized_{_run_id}"):
                try:
                    from database.analysis_steps_manager import finalize_run
                    finalize_run(_run_id, status="cancelled")
                    st.session_state[f"_run_finalized_{_run_id}"] = True
                except Exception as _fin_err:
                    print(f"[AnalysisSteps] finalize_run 오류: {_fin_err}")

            # 세션 저장 후 재시작
            try:
                from auth.session_init import save_work_session
                save_work_session()
            except Exception as e:
                print(f"세션 저장 오류: {e}")

            st.rerun()

        # 건너뛰기 처리
        if skip_clicked:
            # 건너뛴 블록 기록 (선택적)
            if 'skipped_blocks' not in st.session_state:
                st.session_state.skipped_blocks = []
            st.session_state.skipped_blocks.append(next_block_id)

            # analysis_steps 상태 업데이트(있으면)
            try:
                from database.analysis_steps_manager import set_step_status
                step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                sid = step_map.get(next_block_id)
                if sid:
                    set_step_status(sid, "skipped")
            except Exception as _skip_db_err:
                print(f"[AnalysisSteps] skip 업데이트 실패: {_skip_db_err}")

            st.session_state.cot_current_index += 1
            
            # 세션 저장 후 재시작
            try:
                from auth.session_init import save_work_session
                save_work_session()
            except Exception as e:
                print(f"세션 저장 오류: {e}")
            
            st.info(f"{next_block_name} 블록을 건너뛰었습니다.")
            st.rerun()

        if run_clicked:
            analyzer = get_cot_analyzer()
            if analyzer is None:
                st.error("분석기를 초기화할 수 없습니다. 위의 오류 메시지를 확인하세요.")
                st.stop()
            progress_placeholder = st.empty()
            st.session_state.cot_running_block = next_block_id
            # analysis_steps 상태 업데이트(있으면)
            try:
                from database.analysis_steps_manager import set_step_status, save_step_payloads
                step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                sid = step_map.get(next_block_id)
                if sid:
                    set_step_status(sid, "running")
                    _inputs = {
                        "feedback": st.session_state.get("cot_feedback_inputs", {}).get(next_block_id, ""),
                        "spatial_layers": st.session_state.get("block_spatial_selection", {}).get(next_block_id, []),
                    }
                    save_step_payloads(sid, inputs=_inputs, outputs=None)
            except Exception as _run_db_err:
                print(f"[AnalysisSteps] running 업데이트 실패: {_run_db_err}")

            def step_progress(message: str) -> None:
                st.session_state.cot_progress_messages.append(message)
                if len(st.session_state.cot_progress_messages) > 50:
                    st.session_state.cot_progress_messages = st.session_state.cot_progress_messages[-50:]
                progress_placeholder.info(message)

            # 블록별 공간 데이터 컨텍스트 생성
            block_spatial_context = ""
            block_spatial_selection = st.session_state.get('block_spatial_selection', {})
            selected_layers = block_spatial_selection.get(next_block_id, [])
            downloaded_geo_data = st.session_state.get('downloaded_geo_data', {})

            if selected_layers and downloaded_geo_data:
                try:
                    import geopandas as gpd
                    from geo_data_loader import extract_spatial_context_for_ai
                    spatial_parts = []
                    for layer_name in selected_layers:
                        if layer_name in downloaded_geo_data:
                            geo_data = downloaded_geo_data[layer_name]
                            geojson = geo_data.get('geojson', {})
                            features = geojson.get('features', [])
                            if features:
                                gdf = gpd.GeoDataFrame.from_features(features, crs='EPSG:4326')
                                # 레이어 타입 추정
                                layer_type = 'general'
                                if any(kw in layer_name for kw in ['행정', '시군', '읍면', '경계']):
                                    layer_type = 'administrative'
                                elif any(kw in layer_name for kw in ['용도', '지역', '지구']):
                                    layer_type = 'zoning'
                                elif any(kw in layer_name for kw in ['도시계획', '시설']):
                                    layer_type = 'urban_planning'
                                spatial_text = extract_spatial_context_for_ai(gdf, layer_type)
                                spatial_parts.append(f"**{layer_name}**\n{spatial_text}")
                    if spatial_parts:
                        block_spatial_context = "\n\n[공간 데이터 컨텍스트]\n" + "\n\n---\n\n".join(spatial_parts)
                        st.caption(f"📍 {len(spatial_parts)}개 공간 레이어 포함")
                except Exception as e:
                    st.warning(f"공간 데이터 처리 오류: {e}")

            # 피드백과 공간 컨텍스트 결합
            user_feedback = st.session_state.cot_feedback_inputs.get(next_block_id, "").strip()
            combined_feedback = None
            if user_feedback or block_spatial_context:
                parts = []
                if user_feedback:
                    parts.append(user_feedback)
                if block_spatial_context:
                    parts.append(block_spatial_context)
                combined_feedback = "\n\n".join(parts) if parts else None

            # 블록별 RAG 컨텍스트 주입
            doc_rag_system = st.session_state.get('doc_rag_system')
            if doc_rag_system and next_block:
                try:
                    from rag_helper import get_block_relevant_context
                    block_context = get_block_relevant_context(next_block, doc_rag_system, top_k=8)
                    if block_context:
                        # 원본 전체 텍스트는 보존하고 블록 전용 컨텍스트를 주입
                        original_file_text = st.session_state.cot_session["project_info"].get("file_text", "")
                        st.session_state.cot_session["project_info"]["_original_file_text"] = original_file_text
                        st.session_state.cot_session["project_info"]["file_text"] = block_context
                        print(f"[RAG] {next_block_id}: 블록 컨텍스트 주입 ({len(block_context)}자 / 전체 {len(original_file_text)}자)")
                except Exception as _rag_inject_err:
                    print(f"[RAG] 컨텍스트 주입 실패 (전체 문서로 폴백): {_rag_inject_err}")

            try:
                with st.spinner("분석 실행 중..."):
                    step_result = analyzer.run_cot_step(
                        next_block_id,
                        next_block,
                        st.session_state.cot_session,
                        progress_callback=step_progress,
                        step_index=st.session_state.cot_current_index + 1,
                        feedback=combined_feedback
                    )
            finally:
                st.session_state.cot_running_block = None
                # 블록별 컨텍스트 주입 후 원본 텍스트 복원
                if st.session_state.cot_session and "project_info" in st.session_state.cot_session:
                    original = st.session_state.cot_session["project_info"].pop("_original_file_text", None)
                    if original is not None:
                        st.session_state.cot_session["project_info"]["file_text"] = original

            if step_result.get('success'):
                st.session_state.cot_session = step_result['cot_session']
                st.session_state.cot_results[next_block_id] = step_result['analysis']
                analysis_result = step_result['analysis']
                st.session_state.analysis_results[next_block_id] = analysis_result
                
                # Citations 저장
                if step_result.get('all_citations'):
                    st.session_state.cot_citations[next_block_id] = step_result['all_citations']

                # Phase 3: 출처 검증
                _run_rag = st.session_state.get('doc_rag_system')
                if _run_rag and analysis_result:
                    try:
                        from rag_helper import verify_analysis
                        _verifications = verify_analysis(analysis_result, _run_rag, max_claims=6)
                        if _verifications:
                            st.session_state.cot_verifications[next_block_id] = _verifications
                    except Exception as _ver_err:
                        print(f"[Verify] 출처 검증 실패: {_ver_err}")

                # analysis_steps 출력 저장(있으면)
                try:
                    from database.analysis_steps_manager import set_step_status, save_step_payloads
                    step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                    sid = step_map.get(next_block_id)
                    if sid:
                        outp = {
                            "analysis": analysis_result,
                            "citations": st.session_state.get("cot_citations", {}).get(next_block_id),
                            "verifications": st.session_state.get("cot_verifications", {}).get(next_block_id),
                        }
                        save_step_payloads(sid, inputs=None, outputs=outp)
                        set_step_status(sid, "completed")
                except Exception as _save_step_err:
                    print(f"[AnalysisSteps] step 저장 실패: {_save_step_err}")

                # 자동 저장
                project_info = {
                    "project_name": st.session_state.get('project_name', ''),
                    "location": st.session_state.get('location', '')
                }
                save_analysis_result(next_block_id, analysis_result, project_info)

                st.session_state.cot_history = step_result['cot_session'].get('cot_history', st.session_state.cot_history)
                st.session_state.cot_current_index += 1

                # 분석 진행 상태 실시간 저장
                try:
                    from auth.session_init import save_analysis_progress, save_work_session
                    save_analysis_progress(force=True)  # 즉시 저장
                    save_work_session()
                except Exception as e:
                    print(f"세션 저장 오류: {e}")

                st.success(f"{next_block_name} 블록 분석이 완료되었습니다.")
                st.rerun()
            else:
                # 실패 상태 저장(있으면)
                try:
                    from database.analysis_steps_manager import set_step_status
                    step_map = st.session_state.get("analysis_step_id_map", {}) or {}
                    sid = step_map.get(next_block_id)
                    if sid:
                        set_step_status(sid, "failed", error=str(step_result.get("error", ""))[:800])
                except Exception as _fail_step_err:
                    print(f"[AnalysisSteps] failed 업데이트 실패: {_fail_step_err}")
                st.error(f"{next_block_name} 블록 분석 실패: {step_result.get('error', '알 수 없는 오류')}")

    if st.session_state.cot_progress_messages:
        with st.expander("최근 진행 메시지", expanded=False):
            for msg in st.session_state.cot_progress_messages[-10:]:
                st.write(msg)

    if st.session_state.cot_session and st.session_state.cot_plan and st.session_state.cot_current_index >= len(st.session_state.cot_plan):
        st.success("모든 블록에 대한 단계별 분석이 완료되었습니다.")
        # analysis_runs 완료 처리
        _run_id = st.session_state.get("current_analysis_run_id")
        if _run_id and not st.session_state.get(f"_run_finalized_{_run_id}"):
            try:
                from database.analysis_steps_manager import finalize_run
                finalize_run(_run_id, status="completed")
                st.session_state[f"_run_finalized_{_run_id}"] = True
            except Exception as _fin_err:
                print(f"[AnalysisSteps] finalize_run 오류: {_fin_err}")

    # 결과는 cot_results와 analysis_results 둘 다 확인 (동기화 보장)
    analysis_results_state = st.session_state.get('analysis_results', {})
    cot_results_state = st.session_state.get('cot_results', {})

    # 두 저장소를 병합 (cot_results가 최신일 수 있음)
    merged_results = {}
    for block_id in selected_blocks:
        if block_id in analysis_results_state:
            merged_results[block_id] = analysis_results_state[block_id]
        elif block_id in cot_results_state:
            # cot_results에만 있으면 analysis_results에 복사
            merged_results[block_id] = cot_results_state[block_id]
            st.session_state.analysis_results[block_id] = cot_results_state[block_id]

    if merged_results:
        ordered_results = merged_results
        if ordered_results:
            st.subheader("📊 분석 결과 미리보기")
            tab_blocks = list(ordered_results.keys())
            tab_titles = []
            for idx, block_id in enumerate(tab_blocks, start=1):
                block = block_lookup.get(block_id)
                block_name = block.get('name', block_id) if block else block_id
                tab_titles.append(f"{idx}. {block_name}")
            preview_tabs = st.tabs(tab_titles)
            for tab, block_id in zip(preview_tabs, tab_blocks):
                with tab:
                    block = block_lookup.get(block_id)
                    st.markdown("**분석 결과**")
                    render_analysis_result(ordered_results[block_id])

                    # Phase 3: 출처 검증 결과
                    _verif = st.session_state.get('cot_verifications', {}).get(block_id, [])
                    if _verif:
                        with st.expander("🔍 출처 검증", expanded=False):
                            for _v in _verif:
                                _conf = _v.get('confidence', 0.0)
                                if _conf >= 0.15:
                                    _label = "🟢 근거 있음"
                                elif _conf >= 0.05:
                                    _label = "🟡 부분 근거"
                                else:
                                    _label = "🔴 근거 부족"
                                st.markdown(f"**{_label}** (신뢰도: {_conf:.3f})")
                                st.caption(f"주장: {_v.get('claim', '')[:120]}")
                                if _v.get('evidence'):
                                    st.caption(f"근거: {_v['evidence'][:150]}")

    all_blocks_completed = (
        st.session_state.cot_plan
        and len(st.session_state.analysis_results) >= len(st.session_state.cot_plan)
    )
    if all_blocks_completed:
        from datetime import datetime
        import json
        filename = f"analysis_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        ordered_results_for_save = {
            block_id: st.session_state.analysis_results[block_id]
            for block_id in st.session_state.cot_plan
            if block_id in st.session_state.analysis_results
        }
        analysis_record = {
            "project_info": project_info_payload,
            "analysis_results": ordered_results_for_save,
            "analysis_timestamp": datetime.now().isoformat(),
            "cot_history": st.session_state.get('cot_history', []),
            "llm_settings": {
                "temperature": st.session_state.llm_temperature,
                "max_tokens": st.session_state.llm_max_tokens
            }
        }
        st.download_button(
            label="💾 분석 결과 다운로드 (JSON)",
            data=json.dumps(analysis_record, ensure_ascii=False, indent=2),
            file_name=filename,
            mime="application/json",
            use_container_width=True,
        )

with tab_download:
    st.header("결과 다운로드")

    # cot_results와 analysis_results 병합 (간헐적 표시 문제 방지)
    analysis_results = st.session_state.get('analysis_results', {}).copy()
    cot_results = st.session_state.get('cot_results', {})
    for block_id, result in cot_results.items():
        if block_id not in analysis_results:
            analysis_results[block_id] = result

    if not analysis_results:
        st.warning("먼저 분석을 실행해주세요.")
        st.stop()
    
    if analysis_results:
        st.success(f"{len(analysis_results)}개 분석 결과가 준비되었습니다.")
        
        # 현재 사용 중인 AI 모델 정보 표시
        current_provider = get_current_provider()
        provider_config = PROVIDER_CONFIG.get(current_provider, {})
        provider_name = provider_config.get('display_name', current_provider)
        model_name = provider_config.get('model', 'unknown')
        st.caption(f"🤖 현재 사용 중인 AI 모델: {provider_name} ({model_name})")
        
        # Word 문서 생성
        if st.button("Word 문서 생성", type="primary"):
            with st.spinner("Word 문서 생성 중..."):
                doc = create_word_document(project_name, analysis_results)
                
                # 메모리에 직접 바이트 데이터 생성
                import io
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_data = doc_buffer.getvalue()
                
                # 다운로드 버튼 표시
                st.download_button(
                    label="📥 Word 문서 다운로드",
                    data=file_data,
                    file_name=f"{project_name}_분석보고서.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # 개별 결과 다운로드
        st.subheader("개별 분석 결과")
        for block_id, result in analysis_results.items():
            # 블록 이름 찾기
            block_name = "알 수 없음"
            example_blocks = get_example_blocks()
            custom_blocks = load_custom_blocks()
            for block in example_blocks + custom_blocks:
                if block['id'] == block_id:
                    block_name = block['name']
                    break
            
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{block_name}**")
            with col2:
                st.download_button(
                    label="📥 다운로드",
                    data=str(result) if not isinstance(result, (str, bytes)) else result,
                    file_name=f"{block_name}.txt",
                    mime="text/plain",
                    key=f"download_{block_id}"
                )
    else:
        st.info("분석 결과가 없습니다.")

# 페이지 렌더링 완료 후 작업 세션 자동 저장 (3초 스로틀)
try:
    from auth.session_init import auto_save_debounced
    auto_save_debounced(throttle_seconds=3.0)
except Exception as e:
    pass
